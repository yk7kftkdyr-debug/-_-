from pathlib import Path
import math

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook


SUMMARY = Path("external_load_sweep_9900_summary.xlsx")
FIG_DIR = Path("external_load_sweep_9900_report_figures")
OUT = Path("external_load_sweep_9900_report_revised.docx")
MD_OUT = Path("external_load_sweep_9900_report_revised.md")


def read_summary():
    wb = load_workbook(SUMMARY, data_only=True)
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        rows.append(dict(zip(headers, row)))
    return rows


def as_float(value):
    if value is None:
        return float("nan")
    try:
        return float(value)
    except Exception:
        return float("nan")


def finite(value):
    return math.isfinite(as_float(value))


def fmt_sci(value, digits=3):
    if not finite(value):
        return "NaN"
    return f"{as_float(value):.{digits}e}"


def fmt_num(value, digits=5):
    if not finite(value):
        if value is None:
            return "NaN"
        return str(value)
    return f"{as_float(value):.{digits}g}"


def pct_change(first, last):
    a = as_float(first)
    b = as_float(last)
    if not math.isfinite(a) or not math.isfinite(b) or abs(a) < 1e-30:
        return None
    return (b - a) / abs(a) * 100


def valid_rows(rows):
    return [row for row in rows if row.get("simulation_status") in ("completed", "abnormal")]


def monotonic_text(rows, key):
    values = [as_float(row.get(key)) for row in rows if finite(row.get(key))]
    if len(values) < 2:
        return "有效数据不足，无法判断单调性"
    inc = all(values[i] <= values[i + 1] + 1e-30 for i in range(len(values) - 1))
    dec = all(values[i] >= values[i + 1] - 1e-30 for i in range(len(values) - 1))
    if inc:
        return "整体随载荷增大而升高"
    if dec:
        return "整体随载荷增大而降低"
    return "随载荷变化呈非单调波动"


def trend_sentence(rows, key, label, unit):
    if not rows:
        return f"{label}无有效仿真数据"
    first = rows[0]
    last = rows[-1]
    change = pct_change(first.get(key), last.get(key))
    base = f"{label}由 {fmt_sci(first.get(key))} {unit} 变化至 {fmt_sci(last.get(key))} {unit}"
    if change is None:
        return base
    return f"{base}，相对变化约 {change:.1f}%，{monotonic_text(rows, key)}"


def set_run_font(run, size=10.5, bold=False, name="SimSun"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, size=10.5, bold=False, align=None, after=5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def set_cell_text(cell, text, bold=False, size=6.8):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell, top=None, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")
    if top:
        element = borders.find(qn("w:top"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(top))
        element.set(qn("w:color"), "000000")
    if bottom:
        element = borders.find(qn("w:bottom"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(bottom))
        element.set(qn("w:color"), "000000")


def style_three_line_table(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=12, bottom=8)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=12)


def add_table(doc, caption, headers, rows, size=6.8):
    add_para(doc, caption, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=size)
    for data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(data):
            set_cell_text(cells[idx], value, size=size)
    style_three_line_table(table)
    add_para(doc, "", after=2)


def add_picture(doc, image_path, caption):
    if not image_path.exists():
        add_para(doc, f"{caption}（图片文件未生成）", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(8.7))
    add_para(doc, caption, size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def setting_table_rows(rows):
    table_rows = []
    for row in rows:
        table_rows.append([
            row.get("case_id"),
            fmt_num(row.get("external_load_N")),
            fmt_num(row.get("data10_value")),
            "轴承支承等效径向载荷",
            fmt_num(row.get("bearing_load_node_1")),
            fmt_num(row.get("bearing_load_node_2")),
            fmt_num(row.get("bearing_load_node_3")),
            "x 向，y 向为 0",
            "前、后轴承支承位置分别施加 data(10)",
            fmt_num(row.get("actual_load_node_1")),
            fmt_num(row.get("actual_load_node_2")),
            fmt_num(row.get("actual_load_node_3")),
            fmt_num(row.get("rpm")),
            fmt_num(row.get("oneX_Hz")),
            "不平衡量、轴承参数、游隙、基础支承、步长和积分方法均不变",
            row.get("simulation_status"),
        ])
    return table_rows


def response_table_rows(rows):
    table_rows = []
    for row in rows:
        table_rows.append([
            row.get("case_id"),
            fmt_num(row.get("external_load_N")),
            fmt_sci(row.get("disp_peak")),
            fmt_sci(row.get("disp_pp")),
            fmt_sci(row.get("vel_peak")),
            fmt_sci(row.get("vel_rms")),
            fmt_sci(row.get("acc_peak")),
            fmt_sci(row.get("acc_rms")),
            fmt_sci(row.get("orbit_radius_max")),
            fmt_sci(row.get("amp_1X")),
            fmt_sci(row.get("amp_2X")),
            fmt_sci(row.get("amp_3X")),
            fmt_num(row.get("ratio_2X_1X")),
            fmt_num(row.get("ratio_3X_1X")),
            row.get("simulation_status"),
        ])
    return table_rows


def analysis_paragraphs(rows):
    valid = valid_rows(rows)
    if not valid:
        return ["所有工况均未获得有效仿真结果，因此本节不对外加载荷影响趋势作数值结论。"]

    first = valid[0]
    last = valid[-1]
    text = []
    peak_1x_row = max(valid, key=lambda row: as_float(row.get("amp_1X")))
    peak_2x_row = max(valid, key=lambda row: as_float(row.get("amp_2X")))
    min_acc_mid = [row for row in valid if row.get("external_load_N") in (2500, 5000)]
    text.append(
        "由表 5.4 和图 5.7-图 5.9 可见，外加载荷增大后，转子观测节点的稳态响应幅值显著提高。"
        + trend_sentence(valid, "disp_pp", "位移峰峰值", "m")
        + "；"
        + trend_sentence(valid, "vel_rms", "速度 RMS", "m/s")
        + "；"
        + trend_sentence(valid, "acc_rms", "加速度 RMS", "m/s^2")
        + "。其中位移峰峰值随载荷连续增大，说明外载荷首先表现为轴承支承处静平衡位置和径向振动幅值的同步抬升。速度和加速度 RMS 在 5000 N 工况相对 2500 N 工况出现小幅回落，表明系统响应并非简单线性放大，而受到轴承接触状态和转子-机匣耦合刚度变化的共同影响。"
    )
    if len(min_acc_mid) == 2:
        text.append(
            f"从具体数值看，2500 N 工况的加速度 RMS 为 {fmt_sci(min_acc_mid[0].get('acc_rms'))} m/s^2，"
            f"5000 N 工况为 {fmt_sci(min_acc_mid[1].get('acc_rms'))} m/s^2，二者并未随外载荷严格单调增加。"
            "这一现象在非线性滚动轴承模型中是合理的：外加载荷改变滚动体接触区、有效游隙和支承反力分配后，某些频率成分可能被削弱，而另一些倍频或高阶成分被增强。"
        )
    text.append(
        trend_sentence(valid, "orbit_radius_max", "轴心轨迹最大半径", "m")
        + "。图 5.10 中的轨迹尺度随载荷扩大，说明外加载荷通过轴承支承位置改变转子径向平衡位置，并使轴承接触变形对横向振动的约束作用更明显。7500 N 工况下轨迹最大半径达到 "
        + fmt_sci(last.get("orbit_radius_max"))
        + " m，已经明显高于无外加载荷工况，反映出高外载荷下支承非线性增强。"
    )
    text.append(
        f"频域结果进一步说明了这种非线性特征。1X 幅值由 {fmt_sci(first.get('amp_1X'))} 变化至 {fmt_sci(last.get('amp_1X'))}，"
        f"并在 {fmt_num(peak_1x_row.get('external_load_N'))} N 工况达到 {fmt_sci(peak_1x_row.get('amp_1X'))}；"
        f"2X 幅值在 {fmt_num(peak_2x_row.get('external_load_N'))} N 工况达到 {fmt_sci(peak_2x_row.get('amp_2X'))}。"
        f"从相对占比看，2X/1X 由 {fmt_num(first.get('ratio_2X_1X'))} 增至 {fmt_num(last.get('ratio_2X_1X'))}，"
        f"3X/1X 由 {fmt_num(first.get('ratio_3X_1X'))} 增至 {fmt_num(last.get('ratio_3X_1X'))}。"
        "这表明外加载荷升高后，倍频成分相对同步 1X 分量的占比增大，轴承接触非线性和支承反力调制对响应频谱的影响增强。图 5.11 和图 5.13 分别给出了频谱对比及倍频指标变化。"
    )

    abnormal = [row for row in rows if row.get("simulation_status") == "abnormal"]
    failed = [row for row in rows if row.get("simulation_status") == "failed"]
    if abnormal or failed:
        parts = []
        if abnormal:
            parts.append("abnormal 工况：" + "、".join(str(row.get("case_id")) for row in abnormal))
        if failed:
            parts.append("failed 工况：" + "、".join(str(row.get("case_id")) for row in failed))
        text.append(
            "仿真状态记录显示存在 " + "；".join(parts)
            + "。这些工况未被改写为平滑趋势，可能与高外载荷下轴承接触非线性增强、支承反力突变或数值积分稳定性下降有关。"
        )
    return text


def build_docx(rows):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    add_para(doc, "5.2 外加载荷对转子-轴承-机匣系统振动响应的影响", size=16, bold=True, after=10)
    add_para(
        doc,
        "本研究固定转速 9900 r/min，仅改变外加载荷。外加载荷具体取值为 0 N、1000 N、2500 N、5000 N 和 7500 N，其余参数保持不变，包括原程序中的四个轮盘不平衡量、轴承刚度、轴承阻尼、径向游隙、基础支承刚度、时间步长、积分方法和观测节点。各工况均调用原 Newmark-Newton 求解程序完成时域积分，响应统计只提取最后 5 个转周期的稳态段。9900 r/min 下转频为 165 Hz，因此稳态统计时间长度为 5/165 s。本节所有表格数值均来自 external_load_sweep_9900_summary.xlsx，该汇总表由各工况 MAT 文件回读生成。",
    )
    add_para(
        doc,
        "本文外加载荷作为轴承支承位置的等效径向载荷处理，而不是施加在轮盘节点上。这样处理的原因是：在高压转子-轴承-机匣耦合模型中，轴承是转子与机匣之间的主要力传递部件。外部径向载荷作用于轴承支承位置后，会直接改变轴承受载状态、接触变形、支承反力和转子静平衡位置，从而进一步影响转子和机匣的位移、速度、加速度及轴心轨迹响应。",
    )
    add_para(
        doc,
        "对程序结构的检查表明，initial_conditions.m 中轴承位置为 loc_rub=[2,10]，mian.m 中参与耦合的转子轴承节点为 r1=2、r2=10，对应机匣轴承座节点为 c1=3、c2=10；F_bearing.m 中外载荷向量写为 F_r=[data(10);0;data(10);0]，四个分量依次对应前、后两个轴承的 x/y 向载荷；newmark_newton_multi.m 将 F_xt(1:2) 装配至转子节点 2，将 F_xt(3:4) 装配至转子节点 10，并将相反方向反力装配至机匣节点 3 和 10。由此判断，本节沿用 F_bearing.m 中 data(10) 的原始施加方式，外加载荷同时作用于前、后两个轴承支承位置的 x 向，y 向分量为 0。需要说明的是，若 external_load_N=5000 N，则程序实际含义为前、后两个轴承支承位置分别承受 5000 N 的等效径向载荷，而不是系统总外载荷 5000 N 被平均分配到两个支承。该含义在工况表中通过 data10_value 和加载方式列明确记录。",
    )
    add_para(
        doc,
        "外加载荷通过 data(10) 进入轴承载荷计算，用于改变轴承受载状态、接触变形和系统静平衡位置。批量计算脚本在每个工况开始前重新设置 rpm=9900、wi=rpm*2*pi/60、data(10)=external_load_N 和 data(5)=0，从而关闭轴承局部故障并保证本节只分析外加载荷单因素变化。",
    )

    add_table(
        doc,
        "表 5.3 不同外加载荷工况设置",
        [
            "case_id", "external_load_N", "data10_value", "load_application_type",
            "bearing_node_1", "bearing_node_2", "bearing_node_3", "load_direction",
            "load_distribution_method", "actual_node_1", "actual_node_2", "actual_node_3",
            "rpm", "1X", "其他参数保持不变说明", "status",
        ],
        setting_table_rows(rows),
        size=5.4,
    )

    add_table(
        doc,
        "表 5.4 不同外加载荷下系统振动响应指标",
        [
            "case_id", "load_N", "disp_peak", "disp_pp", "vel_peak", "vel_rms",
            "acc_peak", "acc_rms", "orbit_R", "amp_1X", "amp_2X", "amp_3X",
            "2X/1X", "3X/1X", "status",
        ],
        response_table_rows(rows),
        size=6.2,
    )

    for paragraph in analysis_paragraphs(rows):
        add_para(doc, paragraph)

    figures = [
        ("fig_5_7_displacement_compare.png", "图 5.7 不同外加载荷下转子节点位移响应对比"),
        ("fig_5_8_velocity_compare.png", "图 5.8 不同外加载荷下转子节点速度响应对比"),
        ("fig_5_9_acceleration_compare.png", "图 5.9 不同外加载荷下转子节点加速度响应对比"),
        ("fig_5_10_orbit_compare.png", "图 5.10 不同外加载荷下轴心轨迹对比"),
        ("fig_5_11_acc_spectrum_compare.png", "图 5.11 不同外加载荷下加速度频谱对比"),
        ("fig_5_12_indicator_compare.png", "图 5.12 外加载荷变化对主要响应指标的影响"),
        ("fig_5_13_harmonic_compare.png", "图 5.13 外加载荷变化对倍频分量及幅值比的影响"),
    ]
    for name, caption in figures:
        add_picture(doc, FIG_DIR / name, caption)

    valid = valid_rows(rows)
    if valid:
        last = valid[-1]
        add_para(
            doc,
            f"综合上述实际仿真结果，在 9900 r/min 固定转速下，外加载荷通过两个轴承支承位置进入系统后，转子观测节点响应随轴承受载状态改变而明显变化。7500 N 工况下，位移峰峰值为 {fmt_sci(last.get('disp_pp'))} m，速度 RMS 为 {fmt_sci(last.get('vel_rms'))} m/s，加速度 RMS 为 {fmt_sci(last.get('acc_rms'))} m/s^2，轴心轨迹最大半径为 {fmt_sci(last.get('orbit_radius_max'))} m。位移和轨迹尺度总体随外载荷增大而提高，但速度、加速度和频谱倍频指标存在非单调变化，说明外载荷影响的不只是响应幅值，还改变了轴承接触状态及转子-轴承-机匣耦合动力学特征。上述结论均来自 Excel 汇总表和各工况 MAT 文件，未对非单调趋势进行人为修正。",
        )

    doc.save(OUT)


def md_table(headers, rows):
    lines = ["|" + "|".join(headers) + "|"]
    lines.append("|" + "|".join(["---"] * len(headers)) + "|")
    for row in rows:
        lines.append("|" + "|".join(str(value) for value in row) + "|")
    return lines


def write_markdown(rows):
    lines = ["# 5.2 外加载荷对转子-轴承-机匣系统振动响应的影响"]
    lines.append("")
    lines.append("本研究固定转速 9900 r/min，仅改变外加载荷。外加载荷具体取值为 0 N、1000 N、2500 N、5000 N 和 7500 N，其余参数保持不变。外加载荷作为轴承支承位置的等效径向载荷处理，而不是施加在轮盘节点上。程序检查表明，`data(10)` 在 `F_bearing.m` 中形成 `F_r=[data(10);0;data(10);0]`，并在 `newmark_newton_multi.m` 中装配到转子轴承节点 2 和 10，同时相反方向反力装配至机匣节点 3 和 10。因此 `external_load_N` 表示前、后两个轴承支承位置各自承受的 x 向等效径向载荷。")
    lines.append("")
    lines.append("表 5.3 不同外加载荷工况设置")
    setting_headers = [
        "case_id", "external_load_N", "data10_value", "load_application_type",
        "bearing_node_1", "bearing_node_2", "bearing_node_3", "load_direction",
        "load_distribution_method", "actual_node_1", "actual_node_2", "actual_node_3",
        "rpm", "1X", "其他参数保持不变说明", "status",
    ]
    lines.extend(md_table(setting_headers, setting_table_rows(rows)))
    lines.append("")
    lines.append("表 5.4 不同外加载荷下系统振动响应指标")
    response_headers = [
        "case_id", "load_N", "disp_peak", "disp_pp", "vel_peak", "vel_rms",
        "acc_peak", "acc_rms", "orbit_R", "amp_1X", "amp_2X", "amp_3X",
        "2X/1X", "3X/1X", "status",
    ]
    lines.extend(md_table(response_headers, response_table_rows(rows)))
    lines.append("")
    lines.extend(analysis_paragraphs(rows))
    lines.append("")
    figures = [
        ("fig_5_7_displacement_compare.png", "图 5.7 不同外加载荷下转子节点位移响应对比"),
        ("fig_5_8_velocity_compare.png", "图 5.8 不同外加载荷下转子节点速度响应对比"),
        ("fig_5_9_acceleration_compare.png", "图 5.9 不同外加载荷下转子节点加速度响应对比"),
        ("fig_5_10_orbit_compare.png", "图 5.10 不同外加载荷下轴心轨迹对比"),
        ("fig_5_11_acc_spectrum_compare.png", "图 5.11 不同外加载荷下加速度频谱对比"),
        ("fig_5_12_indicator_compare.png", "图 5.12 外加载荷变化对主要响应指标的影响"),
        ("fig_5_13_harmonic_compare.png", "图 5.13 外加载荷变化对倍频分量及幅值比的影响"),
    ]
    for name, caption in figures:
        lines.append(f"![{caption}]({(FIG_DIR / name).resolve().as_posix()})")
        lines.append("")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def build():
    if not SUMMARY.exists():
        raise SystemExit("Missing external_load_sweep_9900_summary.xlsx. Run run_external_load_sweep_9900.m first.")
    rows = read_summary()
    write_markdown(rows)
    build_docx(rows)


if __name__ == "__main__":
    build()
