from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("unbalance_factor_preview.docx")
REF_IMG = Path("docx_reference_extract/media/image7.png")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
TABLE_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, before=0, after=6, line=1.10):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def add_para(doc, text="", size=11, bold=False, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: DARK_BLUE}
    before = {1: 16, 2: 12, 3: 8}
    after = {1: 8, 2: 6, 3: 4}
    p = doc.add_paragraph()
    set_paragraph_format(p, before=before[level], after=after[level])
    r = p.add_run(text)
    set_run_font(r, size=sizes[level], bold=True, color=colors[level])
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=0, line=1.10)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], TABLE_FILL)
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, txt in enumerate(row):
            set_cell_text(cells[i], txt)
    add_para(doc, "", after=2)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, CALLOUT_FILL)
    cell.width = Inches(6.4)
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=2, line=1.10)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    set_paragraph_format(p2, after=0, line=1.10)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    add_para(doc, "", after=2)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, "FAFAFA")
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=0, line=1.0)
    r = p.add_run(code)
    set_run_font(r, size=9.5, name="Consolas")
    add_para(doc, "", after=2)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    add_para(doc, "转子-轴承-机匣系统振动响应单因素分析样稿", size=22, bold=True, color=RGBColor(0, 0, 0), after=4)
    add_para(doc, "因素一：不平衡量对系统振动响应的影响", size=13, color=MUTED, after=12)
    add_para(doc, "基于当前 Newmark 解故障轴承程序包；本稿先展示一个因素的报告组织方式，待确认后可按同一格式扩展外载荷与轴承刚度两部分。", size=10.5, color=MUTED, after=18)

    add_callout(
        doc,
        "本稿定位",
        "这是写作与变量修改方案样稿，不把参考 Word 中的图片作为本程序的计算结果。实际论文图应由各工况仿真后输出的时域图、频谱图、轴心轨迹和幅值指标替换。"
    )

    add_heading(doc, "1 变量修改总览", 1)
    add_para(doc, "当前程序中三类研究因素的入口相对集中。建议每次只改变一个因素，其他参数保持基准值，以免不平衡力、轴承接触力与机匣耦合刚度的影响混在一起。")
    add_table(
        doc,
        ["研究因素", "主要修改位置", "建议变化方式", "保持不变的量"],
        [
            ["不平衡故障", "initial_conditions.m：rpm、U、Famp1、fai；newmark_newton_multi.m：Famp1 施加到 loca 节点", "固定转速与轴承参数，设置 3 组不平衡量，例如 1e-5、5e-5、1e-4 kg·m", "data(5)=0、data(10)、data(6)、kb1/kb2、基础刚度"],
            ["外载荷", "initial_conditions.m：data(10)；F_bearing.m：F_r=[data(10);0;data(10);0]", "设置径向载荷 0、1000、2500、5000 N；如需方向载荷，扩展 F_r 为 Fx/Fy 分量", "U、rpm、data(6)、kb1/kb2"],
            ["轴承刚度", "mian.m：kb1/kb2；F_bearing.m：C_b=data(6)；K_D.m：固定支承模型 Kzx/Kzy", "耦合模型优先改 kb1/kb2；若研究赫兹接触刚度，再单独改 data(6)", "U、rpm、data(10)、游隙 data(9)"],
        ],
        [1.2, 2.0, 2.1, 1.2],
    )

    add_heading(doc, "2 不平衡量单因素方案", 1)
    add_heading(doc, "2.1 程序入口", 2)
    add_para(doc, "程序的实际不平衡激励不是通过 data(5) 控制，而是在 initial_conditions.m 中生成 Famp1，并在 newmark_newton_multi.m 中按 cos/sin 形式施加到转子节点。data(5) 更偏向轴承局部故障类型，为了隔离不平衡影响，建议保持 data(5)=0。")
    add_table(
        doc,
        ["文件", "变量", "含义", "本因素建议"],
        [
            ["initial_conditions.m", "loca=[4,6,8,14]", "不平衡量所在轮盘/节点", "保持不变；先比较同一位置组的激励强度影响"],
            ["initial_conditions.m", "rpm、wi", "转速与角速度", "若复现参考图风格，可固定 rpm=7500；若沿用现程序基准，可固定 rpm=9900"],
            ["initial_conditions.m", "U、Famp1", "不平衡量与不平衡力", "把 ISO 计算替换为直接给定 U_case，Famp1=U*wi^2"],
            ["newmark_newton_multi.m", "Fx/Fy", "不平衡力在 x/y 方向的旋转分量", "不需要改动；只要 Famp1 长度与 loca 一致即可"],
        ],
        [1.5, 1.3, 1.8, 1.8],
    )

    add_heading(doc, "2.2 推荐工况", 2)
    add_table(
        doc,
        ["工况编号", "不平衡量 U / kg·m", "转速 / r·min-1", "观察节点", "输出图"],
        [
            ["U1", "1×10^-5", "7500", "转子轴承节点2、机匣对应节点", "加速度时域、频谱、轴心轨迹、峰峰值"],
            ["U2", "5×10^-5", "7500", "同 U1", "同 U1"],
            ["U3", "1×10^-4", "7500", "同 U1", "同 U1"],
        ],
        [0.8, 1.4, 1.1, 1.6, 1.6],
    )
    add_para(doc, "如果希望沿用当前程序的 9900 r/min 基准，只需把表中转速统一替换为 9900 r/min；三个 U 值仍保持不变，这样能单独考察不平衡量增强对响应的影响。")

    add_heading(doc, "2.3 建议代码修改", 2)
    add_para(doc, "最稳妥的做法是先在 initial_conditions.m 里用直接给定的不平衡量替代 ISO 等级换算段。后续批量仿真时，可把 case_id 改成循环变量或单独写 runner 脚本。")
    add_code_block(
        doc,
        "% ===== 不平衡量单因素设置 =====\n"
        "rpm = 7500;                 % 固定转速，和参考图频谱峰值一致\n"
        "wi = rpm*2*pi/60;           % rad/s\n"
        "U_case = [1e-5, 5e-5, 1e-4];% kg*m\n"
        "case_id = 1;                % 1/2/3 对应三组工况\n"
        "U = U_case(case_id) * ones(1, length(loca));\n"
        "Famp1 = U * wi^2;           % 不平衡力 N\n"
        "fai = [0 0 0 0]/180*pi;     % 初始相位保持一致"
    )
    add_callout(
        doc,
        "口径提醒",
        "若你的论文把 1e-5 kg·m 定义为“总不平衡量”，则建议改为 U = U_total * Mass/sum(Mass)，按轮盘质量比例分配；若定义为“每个轮盘的等效不平衡量”，则使用上面的 ones 写法。报告中必须说明采用哪一种口径。"
    )

    add_heading(doc, "3 报告正文样稿", 1)
    add_para(doc, "为研究不平衡激励强度对转子-轴承-机匣系统振动响应的影响，选取不平衡量 U=1×10^-5 kg·m、5×10^-5 kg·m 和 1×10^-4 kg·m 三种工况，在转速 7500 r/min 下进行 Newmark 时域积分计算。计算中保持轴承接触刚度、转子-机匣耦合刚度、径向游隙、外载荷及阻尼参数不变，仅改变不平衡量，以保证响应差异主要来源于旋转不平衡激励。")
    add_para(doc, "对每组工况，建议提取最后五个转周期的稳态响应，分别绘制转子轴承节点与机匣对应节点的加速度时域图和频谱图。时域响应主要用于观察波形稳定性、幅值放大和包络调制；频域响应重点比较转频、一倍频、二倍频及高频分量的变化。")

    add_heading(doc, "3.1 低不平衡量工况", 2)
    add_para(doc, "当 U=1×10^-5 kg·m 时，转子节点响应应以转频成分为主，时域波形通常表现为较规则的周期振动。若机匣节点幅值明显低于转子节点，说明轴承座及机匣结构对不平衡激励具有较好的传递抑制作用。此时若频谱中高倍频成分较弱，可判断系统整体仍接近线性稳态响应。")

    add_heading(doc, "3.2 中等不平衡量工况", 2)
    add_para(doc, "当 U 增大至 5×10^-5 kg·m 时，不平衡力按 F=U*wi^2 同比例增大，转子节点加速度幅值会明显提高。若时域波形出现轻微包络波动，且频谱中二倍频或局部高频分量增强，说明轴承非线性接触、转子-机匣耦合或支承刚度非线性已经开始参与响应。机匣节点若同步出现二倍频成分，表明振动能量从转子侧向机匣侧的传递增强。")

    add_heading(doc, "3.3 高不平衡量工况", 2)
    add_para(doc, "当 U=1×10^-4 kg·m 时，系统可能进入强激励状态。此时应重点观察三个现象：第一，时域峰值是否接近或超过径向游隙 data(9)；第二，轴心轨迹是否由近椭圆形转变为畸变轨迹；第三，频谱中二倍频、三倍频及宽频成分是否显著增强。若这些现象同时出现，可认为不平衡激励已经诱发明显的非线性振动响应。")

    add_heading(doc, "4 图表组织方式", 1)
    add_para(doc, "建议每个不平衡量工况用一组四联图呈现：左上为转子节点加速度时域，右上为机匣节点加速度时域，左下为转子节点频谱，右下为机匣节点频谱。图题可沿用“图4.x 不平衡量为 U 下转子及机匣节点振动响应”的格式。")
    if REF_IMG.exists():
        add_para(doc, "参考图式示例：频谱图应标出转频及倍频峰值，并用虚线圈或箭头强调高频成分变化。下图仅作为排版与标注风格参考，不作为本程序仿真结果。", color=MUTED)
        doc.add_picture(str(REF_IMG), width=Inches(4.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_para(doc, "图A 参考 Word 中的频谱标注风格示例", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    add_heading(doc, "5 小结写法", 1)
    add_para(doc, "不平衡量增大时，转子节点响应通常表现为幅值递增、波形由规则周期振动向包络调制发展、频谱由单一转频向倍频及宽频成分扩展。机匣节点响应若随之增强，则说明振动能量经轴承支承与机匣耦合路径向外传递，隔振能力下降。最终报告建议用峰值、均方根值、主频幅值和倍频幅值比四个指标进行定量对比，避免只做定性描述。")

    doc.save(OUT)


if __name__ == "__main__":
    build()
