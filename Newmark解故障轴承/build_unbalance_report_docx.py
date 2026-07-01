from pathlib import Path
import math

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook


SUMMARY = Path("unbalance_sweep_9900_summary.xlsx")
FIG_DIR = Path("unbalance_sweep_9900_figures")
OUT = Path("unbalance_sweep_9900_report.docx")
MD_OUT = Path("unbalance_sweep_9900_report.md")


def read_summary():
    wb = load_workbook(SUMMARY, data_only=True)
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        rows.append(dict(zip(headers, row)))
    return rows


def fmt_sci(x, digits=3):
    if x is None:
        return ""
    try:
        if math.isnan(float(x)):
            return "NaN"
    except Exception:
        pass
    return f"{float(x):.{digits}e}"


def fmt_num(x, digits=4):
    if x is None:
        return ""
    try:
        if math.isnan(float(x)):
            return "NaN"
    except Exception:
        pass
    return f"{float(x):.{digits}g}"


def pct(a, b):
    a = float(a)
    b = float(b)
    if abs(a) < 1e-30:
        return None
    return (b - a) / abs(a) * 100


def first_nonzero(rows, key):
    for row in rows:
        if abs(float(row[key])) > 1e-30:
            return row
    return rows[0]


def trend_text(rows, key, name, unit):
    start = first_nonzero(rows, key)
    end = rows[-1]
    change = pct(start[key], end[key])
    monotonic = all(float(rows[i][key]) <= float(rows[i + 1][key]) + 1e-30 for i in range(len(rows) - 1))
    if change is None:
        change_txt = f"{name} 在基准非零工况后达到 {fmt_sci(end[key])} {unit}"
    else:
        change_txt = f"{name} 由 {fmt_sci(start[key])} {unit} 增至 {fmt_sci(end[key])} {unit}，增幅约为 {change:.1f}%"
    if monotonic:
        return change_txt + "，整体呈随不平衡量增大的上升趋势"
    return change_txt + "，但中间工况存在非单调波动"


def set_run_font(run, size=12, bold=False, name="SimSun"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, size=12, bold=False, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell, top=None, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")
    if top:
        element = borders.find(qn("w:top"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(top))
        element.set(qn("w:color"), "000000")
    if bottom:
        element = borders.find(qn("w:bottom"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(bottom))
        element.set(qn("w:color"), "000000")


def style_three_line_table(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=12, bottom=8)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=12)


def add_table(doc, caption, headers, data_rows, widths=None, size=8):
    add_para(doc, caption, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=size)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for data in data_rows:
        cells = table.add_row().cells
        for i, value in enumerate(data):
            set_cell_text(cells[i], value, size=size)
            if widths:
                cells[i].width = Inches(widths[i])
    style_three_line_table(table)
    add_para(doc, "", after=2)


def add_picture(doc, image_path, caption):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.7))
    add_para(doc, caption, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def build():
    rows = read_summary()
    write_markdown(rows)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    add_para(doc, "5.1 不平衡量对转子—轴承—机匣系统振动响应的影响", size=16, bold=True, after=10)

    scales = [row["unbalance_scale"] for row in rows]
    first = first_nonzero(rows, "disp_pp")
    last = rows[-1]
    disp_sentence = trend_text(rows, "disp_pp", "转子节点位移峰峰值", "m")
    vel_sentence = trend_text(rows, "vel_rms", "速度 RMS", "m/s")
    acc_sentence = trend_text(rows, "acc_rms", "加速度 RMS", "m/s^2")
    orbit_sentence = trend_text(rows, "orbit_radius_max", "轴心轨迹最大半径", "m")
    one_x_dominant = all(float(row["amp_1X"]) >= float(row["amp_2X"]) and float(row["amp_1X"]) >= float(row["amp_3X"]) for row in rows)
    ratio2_change = pct(first["ratio_2X_1X"], last["ratio_2X_1X"])
    ratio3_change = pct(first["ratio_3X_1X"], last["ratio_3X_1X"])
    case_change = pct(first["case_acc_rms"], last["case_acc_rms"]) if "case_acc_rms" in first else None

    add_para(doc, "本研究固定转速为 9900 r/min，对应转频 1X 为 165 Hz。计算过程中保持外载荷、轴承刚度、轴承阻尼、径向游隙、基础支承刚度及其他结构参数不变，仅改变四个轮盘的不平衡量。四个轮盘不平衡位置保持原模型设置，即 loca=[4, 6, 8, 14]，各轮盘不平衡量按相同比例系数整体放大或减小，比例系数依次取 0、0.5、1、2 和 5，从而保证不平衡分布位置和相对比例与原模型一致。")
    add_para(doc, "第 i 个轮盘的不平衡量记为 U_i，其对应的不平衡离心力幅值为：")
    add_para(doc, "F_i = U_i · wi^2", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "其中，wi 为转子角速度。由于本节固定 rpm=9900 r/min，不同工况下不平衡力幅值的变化仅由 U_i 的比例变化决定。每个工况均重新计算四个轮盘的不平衡力幅值，并保证 Famp1 与 loca 均为四元素向量，避免只改变单个轮盘而破坏原模型的不平衡分布特征。")
    add_para(doc, "对每组不平衡量工况，提取最后五个转周期的稳态响应，输出转子观测节点和机匣对应节点的位移、速度、加速度、轴心轨迹及频谱结果，并计算位移峰值、位移峰峰值、速度 RMS、加速度 RMS、轴心轨迹最大半径、1X 幅值、2X/1X 幅值比和 3X/1X 幅值比等指标。")

    setting_rows = []
    for i, row in enumerate(rows, start=1):
        setting_rows.append([
            f"U{i}",
            fmt_num(row["unbalance_scale"]),
            fmt_sci(row["U1"]), fmt_sci(row["U2"]), fmt_sci(row["U3"]), fmt_sci(row["U4"]),
            fmt_sci(row["Famp1_1"]), fmt_sci(row["Famp1_2"]), fmt_sci(row["Famp1_3"]), fmt_sci(row["Famp1_4"]),
            "9900", "165",
        ])
    add_table(
        doc,
        "表 5.1 不同不平衡量工况设置",
        ["工况", "scale", "U1", "U2", "U3", "U4", "F1", "F2", "F3", "F4", "rpm", "1X/Hz"],
        setting_rows,
        size=6.5,
    )

    response_rows = []
    for row in rows:
        response_rows.append([
            fmt_num(row["unbalance_scale"]),
            fmt_sci(row["disp_peak"]), fmt_sci(row["disp_pp"]),
            fmt_sci(row["vel_peak"]), fmt_sci(row["vel_rms"]),
            fmt_sci(row["acc_peak"]), fmt_sci(row["acc_rms"]),
            fmt_sci(row["orbit_radius_max"]),
            fmt_sci(row["amp_1X"]), fmt_sci(row["amp_2X"]), fmt_sci(row["amp_3X"]),
            fmt_num(row["ratio_2X_1X"]), fmt_num(row["ratio_3X_1X"]),
        ])
    add_table(
        doc,
        "表 5.2 不同不平衡量下系统振动响应指标",
        ["scale", "disp_peak", "disp_pp", "vel_peak", "vel_rms", "acc_peak", "acc_rms", "orbit_R", "amp_1X", "amp_2X", "amp_3X", "2X/1X", "3X/1X"],
        response_rows,
        size=6.2,
    )

    add_para(doc, f"由表 5.2 可见，{disp_sentence}；{vel_sentence}；{acc_sentence}。与位移和速度相比，加速度响应直接反映高频动态分量，因此当不平衡量增强时，其变化更容易揭示轴承接触和转子—机匣耦合带来的动态放大。{orbit_sentence}，说明不平衡激励增强会扩大轴心运动范围；若图 5.4 中高比例工况轨迹出现局部弯折或厚化，则表明系统响应已不再是简单的线性椭圆轨迹。")

    add_picture(doc, FIG_DIR / "fig_5_1_displacement_compare.png", "图 5.1 不同不平衡量下转子节点位移响应对比")
    add_picture(doc, FIG_DIR / "fig_5_2_velocity_compare.png", "图 5.2 不同不平衡量下转子节点速度响应对比")
    add_picture(doc, FIG_DIR / "fig_5_3_acceleration_compare.png", "图 5.3 不同不平衡量下转子节点加速度响应对比")
    add_picture(doc, FIG_DIR / "fig_5_4_orbit_compare.png", "图 5.4 不同不平衡量下轴心轨迹对比")
    add_picture(doc, FIG_DIR / "fig_5_5_acc_spectrum_compare.png", "图 5.5 不同不平衡量下加速度频谱对比")
    add_picture(doc, FIG_DIR / "fig_5_6_indicator_compare.png", "图 5.6 不平衡量变化对系统响应指标的影响")

    if one_x_dominant:
        spectrum_text = "各工况下 1X 幅值均高于 2X 和 3X 幅值，说明系统响应仍主要由同步不平衡离心激励控制。"
    else:
        spectrum_text = "部分工况中 2X 或 3X 幅值接近甚至超过 1X 幅值，说明系统响应中已经出现较明显的非同步或非线性成分。"
    ratio_bits = []
    if ratio2_change is not None:
        ratio_bits.append(f"2X/1X 幅值比由 {fmt_num(first['ratio_2X_1X'])} 变化至 {fmt_num(last['ratio_2X_1X'])}，相对变化约为 {ratio2_change:.1f}%")
    if ratio3_change is not None:
        ratio_bits.append(f"3X/1X 幅值比由 {fmt_num(first['ratio_3X_1X'])} 变化至 {fmt_num(last['ratio_3X_1X'])}，相对变化约为 {ratio3_change:.1f}%")
    add_para(doc, spectrum_text + " " + "；".join(ratio_bits) + "。若高倍频比值随 scale 增大而升高，则说明轴承非线性接触及转子—机匣耦合作用逐渐增强；若高倍频比值始终较小，则说明系统在本参数范围内仍以稳定同步响应为主。")

    if case_change is None:
        case_text = "机匣节点响应随转子响应变化而发生同步变化。"
    else:
        case_text = f"机匣节点加速度 RMS 由 {fmt_sci(first['case_acc_rms'])} m/s^2 变化至 {fmt_sci(last['case_acc_rms'])} m/s^2，变化幅度约为 {case_change:.1f}%。"
    add_para(doc, case_text + " 这表明转子侧由不平衡激励产生的振动能量能够通过轴承支承路径传递至机匣结构，且传递强度随不平衡激励增强而提高。")

    add_para(doc, f"综合上述结果，在固定 9900 r/min 条件下，四个轮盘不平衡量整体放大后，同步离心激励随 U_i 增大而增强，转子节点位移、速度和加速度响应总体上升。频谱结果显示，{'1X 成分保持主导，系统主要受同步不平衡激励控制' if one_x_dominant else '高倍频成分已明显参与响应，系统非线性特征增强'}；同时，机匣节点响应与转子响应共同增大，说明振动能量经轴承支承向机匣传递。不平衡量不仅影响响应幅值，也会改变轴心轨迹尺度和频谱倍频结构，是转子—轴承—机匣系统振动响应分析中的关键激励参数。")

    doc.save(OUT)


def write_markdown(rows):
    first = first_nonzero(rows, "disp_pp")
    last = rows[-1]
    one_x_dominant = all(float(row["amp_1X"]) >= float(row["amp_2X"]) and float(row["amp_1X"]) >= float(row["amp_3X"]) for row in rows)
    disp_sentence = trend_text(rows, "disp_pp", "位移峰峰值", "m")
    vel_sentence = trend_text(rows, "vel_rms", "速度 RMS", "m/s")
    acc_sentence = trend_text(rows, "acc_rms", "加速度 RMS", "m/s^2")
    orbit_sentence = trend_text(rows, "orbit_radius_max", "轴心轨迹最大半径", "m")
    ratio2_change = pct(first["ratio_2X_1X"], last["ratio_2X_1X"])
    ratio3_change = pct(first["ratio_3X_1X"], last["ratio_3X_1X"])
    case_change = pct(first["case_acc_rms"], last["case_acc_rms"])
    lines = []
    lines.append("# 不平衡激励对转子-轴承-机匣系统振动响应的影响\n")
    lines.append("本文固定转速为 9900 r/min，转频 1X=165 Hz，保持轴承刚度、阻尼、外载荷、径向游隙、基础支承刚度和其他结构参数不变，仅改变四个轮盘的不平衡量。四个不平衡量按相同比例系数整体放大或减小，比例系数为 `[0, 0.5, 1, 2, 5]`。\n")
    lines.append("不平衡量与离心力幅值的计算关系为：`U_i = m_i e_i`，`F_i = U_i * wi^2`。由于本节固定转速，工况间不平衡力幅值变化完全由 `U_i` 的比例放大引起。\n")
    lines.append(f"从第一个非零不平衡工况到强不平衡工况，{disp_sentence}；{vel_sentence}；{acc_sentence}；{orbit_sentence}。\n")
    if one_x_dominant:
        lines.append("频谱中 1X 成分在各工况下均高于 2X 与 3X 成分，说明系统响应主要受同步不平衡激励控制。")
    else:
        lines.append("部分工况中 2X 或 3X 成分接近或超过 1X 成分，说明响应中已经出现更明显的非同步或非线性成分。")
    ratio_text = []
    if ratio2_change is not None:
        ratio_text.append(f"2X/1X 幅值比由 {fmt_num(first['ratio_2X_1X'])} 变化至 {fmt_num(last['ratio_2X_1X'])}，相对变化约为 {ratio2_change:.1f}%")
    if ratio3_change is not None:
        ratio_text.append(f"3X/1X 幅值比由 {fmt_num(first['ratio_3X_1X'])} 变化至 {fmt_num(last['ratio_3X_1X'])}，相对变化约为 {ratio3_change:.1f}%")
    if ratio_text:
        lines.append(" " + "；".join(ratio_text) + "。")
    if case_change is not None:
        lines.append(f"机匣节点加速度 RMS 由 {fmt_sci(first['case_acc_rms'])} m/s^2 变化至 {fmt_sci(last['case_acc_rms'])} m/s^2，变化幅度约为 {case_change:.1f}%，表明振动能量经轴承支承路径向机匣传递。\n")
    lines.append("## 主要图件\n")
    figs = [
        ("图5.1 不同不平衡量下转子节点位移响应对比", "fig_5_1_displacement_compare.png"),
        ("图5.2 不同不平衡量下转子节点速度响应对比", "fig_5_2_velocity_compare.png"),
        ("图5.3 不同不平衡量下转子节点加速度响应对比", "fig_5_3_acceleration_compare.png"),
        ("图5.4 不同不平衡量下轴心轨迹对比", "fig_5_4_orbit_compare.png"),
        ("图5.5 不同不平衡量下加速度频谱对比", "fig_5_5_acc_spectrum_compare.png"),
        ("图5.6 不平衡量变化对系统响应指标的影响", "fig_5_6_indicator_compare.png"),
    ]
    for caption, name in figs:
        path = (FIG_DIR / name).resolve().as_posix()
        lines.append(f"![{caption}]({path})\n")
    lines.append("指标汇总见 `unbalance_sweep_9900_summary.xlsx`。\n")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    if not SUMMARY.exists():
        raise SystemExit("Missing unbalance_sweep_9900_summary.xlsx. Run run_unbalance_sweep_9900.m first.")
    build()
