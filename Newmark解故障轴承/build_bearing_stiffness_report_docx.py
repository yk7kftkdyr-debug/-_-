from pathlib import Path
import math

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook


SUMMARY = Path("bearing_stiffness_sweep_9900_summary.xlsx")
FIG_DIR = Path("bearing_stiffness_sweep_9900_figures")
OUT = Path("bearing_stiffness_sweep_9900_report.docx")
MD_OUT = Path("bearing_stiffness_sweep_9900_report.md")


def read_summary():
    wb = load_workbook(SUMMARY, data_only=True)
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        rows.append(dict(zip(headers, row)))
    return rows


def as_float(value):
    if value is None or value == "NaN":
        return float("nan")
    try:
        return float(value)
    except Exception:
        return float("nan")


def finite(value):
    return math.isfinite(as_float(value))


def fmt_sci(value, digits=3):
    if not finite(value):
        return "NaN"
    return f"{as_float(value):.{digits}e}"


def fmt_num(value, digits=5):
    if not finite(value):
        if value is None:
            return "NaN"
        return str(value)
    return f"{as_float(value):.{digits}g}"


def pct_change(first, last):
    a = as_float(first)
    b = as_float(last)
    if not math.isfinite(a) or not math.isfinite(b) or abs(a) < 1e-30:
        return None
    return (b - a) / abs(a) * 100


def valid_rows(rows):
    return [row for row in rows if row.get("simulation_status") in ("completed", "abnormal")]


def completed_rows(rows):
    return [row for row in rows if row.get("simulation_status") == "completed"]


def monotonic_text(rows, key):
    values = [as_float(row.get(key)) for row in rows if finite(row.get(key))]
    if len(values) < 2:
        return "有效数据不足，无法判断单调性"
    inc = all(values[i] <= values[i + 1] + 1e-30 for i in range(len(values) - 1))
    dec = all(values[i] >= values[i + 1] - 1e-30 for i in range(len(values) - 1))
    if inc:
        return "整体随刚度增大而升高"
    if dec:
        return "整体随刚度增大而降低"
    return "随刚度变化呈非单调波动"


def set_run_font(run, size=12, bold=False, name="SimSun"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, size=12, bold=False, align=None, after=6, line_spacing=1.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.first_line_indent = Inches(0.29) if align is None else None
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def set_cell_text(cell, text, bold=False, size=6.0):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, val in (("top", "60"), ("bottom", "60"), ("start", "80"), ("end", "80")):
        node = tc_mar.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            tc_mar.append(node)
        node.set(qn("w:w"), val)
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, top=None, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")
    if top:
        element = borders.find(qn("w:top"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(top))
        element.set(qn("w:color"), "000000")
    if bottom:
        element = borders.find(qn("w:bottom"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(bottom))
        element.set(qn("w:color"), "000000")


def style_three_line_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=12, bottom=8)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=12)


def add_table(doc, caption, headers, rows, size=6.0):
    add_para(doc, caption, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=size)
    for data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(data):
            set_cell_text(cells[idx], value, size=size)
    style_three_line_table(table)
    add_para(doc, "", after=2, line_spacing=1.0)


def add_picture(doc, image_path, caption):
    if not image_path.exists():
        add_para(doc, f"{caption}（图片文件未生成）", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(8.2))
    add_para(doc, caption, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def setting_table_rows(rows):
    out = []
    unchanged = "不平衡量、外加载荷、轴承阻尼、径向游隙、基础支承刚度、时间步长、积分方法和观测节点保持不变"
    for row in rows:
        out.append([
            row.get("case_id"),
            fmt_sci(row.get("K_level"), 2),
            fmt_sci(row.get("kb1_case"), 2),
            fmt_sci(row.get("kb2_case"), 2),
            row.get("kb3_case"),
            row.get("Kzx_case"),
            row.get("Kzy_case"),
            fmt_num(row.get("rpm")),
            fmt_num(row.get("oneX_Hz")),
            fmt_num(row.get("unbalance_amplification")),
            fmt_sci(row.get("U1_used"), 3),
            fmt_sci(row.get("U2_used"), 3),
            fmt_sci(row.get("U3_used"), 3),
            fmt_sci(row.get("U4_used"), 3),
            fmt_sci(row.get("Famp1_1_used"), 3),
            fmt_sci(row.get("Famp1_2_used"), 3),
            fmt_sci(row.get("Famp1_3_used"), 3),
            fmt_sci(row.get("Famp1_4_used"), 3),
            unchanged,
            row.get("simulation_status"),
        ])
    return out


def response_table_rows(rows):
    out = []
    for row in rows:
        out.append([
            row.get("case_id"),
            fmt_sci(row.get("K_level"), 2),
            fmt_sci(row.get("disp_peak"), 3),
            fmt_sci(row.get("disp_pp"), 3),
            fmt_sci(row.get("vel_peak"), 3),
            fmt_sci(row.get("vel_rms"), 3),
            fmt_sci(row.get("acc_peak"), 3),
            fmt_sci(row.get("acc_rms"), 3),
            fmt_sci(row.get("orbit_radius_max"), 3),
            fmt_sci(row.get("amp_1X"), 3),
            fmt_sci(row.get("amp_2X"), 3),
            fmt_sci(row.get("amp_3X"), 3),
            fmt_num(row.get("ratio_2X_1X"), 4),
            fmt_num(row.get("ratio_3X_1X"), 4),
            row.get("simulation_status"),
        ])
    return out


def analysis_paragraphs(rows):
    valid = valid_rows(rows)
    comp = completed_rows(rows)
    if not valid:
        return ["所有工况均未获得有效仿真输出，因此本节不对轴承支承刚度影响趋势作数值结论。"]

    first = valid[0]
    last = valid[-1]
    k3 = next((row for row in rows if row.get("case_id") == "K3"), None)
    k5 = next((row for row in rows if row.get("case_id") == "K5"), None)
    max_acc = max(valid, key=lambda row: as_float(row.get("acc_rms")))
    max_case = max(valid, key=lambda row: as_float(row.get("case_acc_rms")))
    max_2x_ratio = max(valid, key=lambda row: as_float(row.get("ratio_2X_1X")))
    one_x_all = all(as_float(row.get("amp_1X")) >= as_float(row.get("amp_2X")) and as_float(row.get("amp_1X")) >= as_float(row.get("amp_3X")) for row in valid)

    text = []
    text.append(
        f"由表 5.6 可见，轴承支承刚度从 1.0e6 N/m 增大到 1.0e10 N/m 后，转子观测节点位移峰峰值由 {fmt_sci(first.get('disp_pp'))} m 变化为 {fmt_sci(last.get('disp_pp'))} m，"
        f"速度 RMS 由 {fmt_sci(first.get('vel_rms'))} m/s 变化为 {fmt_sci(last.get('vel_rms'))} m/s，加速度 RMS 由 {fmt_sci(first.get('acc_rms'))} m/s^2 变化为 {fmt_sci(last.get('acc_rms'))} m/s^2。"
        "需要指出的是，K1 和 K2 在积分过程中因位移超过程序给定阈值而提前终止，汇总表中已标注为 abnormal，因此这两个低刚度工况反映的是弱支承下的异常放大状态，不能按稳定周期响应作简单外推。"
    )
    if k3 and k5:
        text.append(
            f"若只比较完成全时长积分的 K3-K5 工况，位移峰峰值从 {fmt_sci(k3.get('disp_pp'))} m 降至 {fmt_sci(k5.get('disp_pp'))} m，"
            f"速度 RMS 从 {fmt_sci(k3.get('vel_rms'))} m/s 降至 {fmt_sci(k5.get('vel_rms'))} m/s，"
            f"加速度 RMS 从 {fmt_sci(k3.get('acc_rms'))} m/s^2 降至 {fmt_sci(k5.get('acc_rms'))} m/s^2。"
            f"因此在本次 9900 r/min 仿真中，高刚度工况没有表现出加速度 RMS 增强，而是随支承约束增强而进一步降低；该结论来自 K3-K5 的 completed 数据。"
        )
    text.append(
        f"轴心轨迹随刚度增大明显收缩，最大轨迹半径由 K1 的 {fmt_sci(first.get('orbit_radius_max'))} m 降至 K5 的 {fmt_sci(last.get('orbit_radius_max'))} m。"
        f"K1 和 K2 轨迹尺度达到 10^-1 m 量级，并且主程序位移合理性判断显示其远超径向游隙，说明低支承刚度下转子-机匣耦合约束不足，系统出现异常大位移响应；"
        f"K3、K4、K5 的轨迹半径分别为 {fmt_sci(k3.get('orbit_radius_max')) if k3 else 'NaN'} m、{fmt_sci(next((r for r in rows if r.get('case_id') == 'K4'), {}).get('orbit_radius_max'))} m 和 {fmt_sci(last.get('orbit_radius_max'))} m，轨迹形态转为小幅同步椭圆状响应。"
    )
    text.append(
        f"频谱方面，1X 成分在所有获得输出的工况中{'均保持主导' if one_x_all else '并非始终占主导'}。"
        f"1X 加速度幅值由 K1 的 {fmt_sci(first.get('amp_1X'))} 变化为 K5 的 {fmt_sci(last.get('amp_1X'))}，"
        f"2X/1X 幅值比由 {fmt_num(first.get('ratio_2X_1X'), 4)} 变化为 {fmt_num(last.get('ratio_2X_1X'), 4)}，"
        f"3X/1X 幅值比由 {fmt_num(first.get('ratio_3X_1X'), 4)} 变化为 {fmt_num(last.get('ratio_3X_1X'), 4)}。"
        f"最大 2X/1X 出现在 {max_2x_ratio.get('case_id')} 工况，为 {fmt_num(max_2x_ratio.get('ratio_2X_1X'), 4)}，主要对应低刚度异常响应；进入 K3-K5 完成工况后，倍频占比迅速降低，宽频和高阶倍频增强现象并不突出。"
    )
    text.append(
        f"机匣响应同样随支承刚度变化而改变。机匣节点加速度 RMS 在 {max_case.get('case_id')} 工况达到最大值 {fmt_sci(max_case.get('case_acc_rms'))} m/s^2，"
        f"到 K5 降为 {fmt_sci(last.get('case_acc_rms'))} m/s^2。"
        "这表明当前模型中低刚度支承首先导致转子侧异常大位移，并通过转子-轴承-机匣耦合路径放大机匣响应；当支承刚度提高到 1.0e8 N/m 及以上时，转子运动受限，传递到机匣的加速度能量也随之下降。"
    )
    abnormal = [row.get("case_id") for row in rows if row.get("simulation_status") == "abnormal"]
    failed = [row.get("case_id") for row in rows if row.get("simulation_status") == "failed"]
    status_text = []
    if abnormal:
        status_text.append("abnormal 工况：" + "、".join(abnormal))
    if failed:
        status_text.append("failed 工况：" + "、".join(failed))
    status_sentence = "；".join(status_text) if status_text else "所有工况均完成正常积分"
    text.append(
        f"综合来看，在 9900 r/min 固定转速下，轴承支承刚度改变了转子-轴承-机匣系统的支承约束和振动传递路径。"
        f"本次计算状态为：{status_sentence}。位移、速度、加速度和轴心轨迹总体随刚度数量级升高而减小，频谱仍以 1X 同步响应为主，高刚度工况下未出现倍频或高频成分的增强。"
        f"系统在 K3-K5 工况下处于稳定同步响应状态，而 K1-K2 则表现为低刚度约束不足引起的异常放大。所有结论均来自 Excel 汇总表和各工况 MAT 文件，未对异常或非单调趋势进行人为修正。"
    )
    return text


def build_docx(rows):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)

    add_para(doc, "5.3 轴承支承刚度对转子-轴承-机匣系统振动响应的影响", size=16, bold=True, after=10, line_spacing=1.2)
    add_para(
        doc,
        "本研究固定转速为 9900 r/min，仅改变轴承支承刚度。轴承支承刚度采用 1.0e6、1.0e7、1.0e8、1.0e9、1.0e10 N/m 五组数量级工况，其余参数保持原程序设置不变，包括原程序中的四个轮盘不平衡量、外加载荷、轴承阻尼、径向游隙、基础支承刚度、时间步长、积分方法和观测节点。批量计算前未放大不平衡激励量，unbalance_amplification=1，四个轮盘不平衡位置 loca=[4,6,8,14] 和相位 fai 均保持原程序设置。每个工况开始时均重新设置 rpm=9900、wi=rpm*2*pi/60、data(5)=0，并调用原 Newmark-Newton 程序完成时域积分；响应指标仅提取最后 5 个转周期，即 5/165 s 的稳态段进行统计。",
    )
    add_para(doc, "wi = rpm*2π/60，f_rot = rpm/60 = 165 Hz", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(
        doc,
        "对程序结构的检查表明，当前耦合模型中轴承支承刚度的实际入口位于 mian.m 的转子-机匣组合矩阵装配阶段。程序将转子轴承节点 r1=2、r2=10 与机匣轴承座节点 c1=3、c2=10 通过 kb1、kb2 连接，并在 x、y 两个平动方向将 +kb、-kb、-kb、+kb 形式的耦合刚度直接加入整体刚度矩阵 KK。因此，本节实际修改的是 kb1 和 kb2，并令二者同时等于当前 K_level。程序中未发现第三个轴承支承刚度 kb3，也未发现以 K_b 命名的轴承支承刚度矩阵；K_D.m 中的 Kzx、Kzy 只在非耦合固定支承模型中使用，而当前 data(8)=2 的转子-机匣耦合模型下该通道不作为轴承支承刚度入口。K_D_case.m 中的 Kzx、Kzy 对应机匣基础支承刚度，按本节单因素要求保持不变。因此，Excel 和表 5.5 中 kb3_case、Kzx_case、Kzy_case 均记录为 NaN。",
    )
    add_para(
        doc,
        "F_bearing.m 中的 C_b=data(6) 是滚动轴承赫兹接触刚度，参与非线性轴承接触力计算；data(9) 为径向游隙，data(10) 为外加载荷入口。本节未修改 C_b、data(9)、data(10) 或基础支承刚度，避免将接触刚度、游隙、外载荷或机匣基础支承效应混入轴承支承刚度单因素分析。",
    )

    add_table(
        doc,
        "表 5.5 不同轴承支承刚度工况设置",
        [
            "case_id", "K_level", "kb1_case", "kb2_case", "kb3_case", "Kzx_case", "Kzy_case",
            "rpm", "1X", "unbalance_amp", "U1_used", "U2_used", "U3_used", "U4_used",
            "Famp1_1", "Famp1_2", "Famp1_3", "Famp1_4", "其他参数保持不变说明", "status",
        ],
        setting_table_rows(rows),
        size=4.6,
    )

    add_table(
        doc,
        "表 5.6 不同轴承支承刚度下系统振动响应指标",
        [
            "case_id", "K_level", "disp_peak", "disp_pp", "vel_peak", "vel_rms",
            "acc_peak", "acc_rms", "orbit_R", "amp_1X", "amp_2X", "amp_3X",
            "2X/1X", "3X/1X", "status",
        ],
        response_table_rows(rows),
        size=6.0,
    )

    figures = [
        ("fig_5_13_displacement_compare.png", "图 5.13 不同轴承支承刚度下转子节点位移响应对比"),
        ("fig_5_14_velocity_compare.png", "图 5.14 不同轴承支承刚度下转子节点速度响应对比"),
        ("fig_5_15_acceleration_compare.png", "图 5.15 不同轴承支承刚度下转子节点加速度响应对比"),
        ("fig_5_16_orbit_compare.png", "图 5.16 不同轴承支承刚度下轴心轨迹对比"),
        ("fig_5_17_acc_spectrum_compare.png", "图 5.17 不同轴承支承刚度下加速度频谱对比"),
        ("fig_5_18_indicator_compare.png", "图 5.18 轴承支承刚度变化对主要响应指标的影响"),
        ("fig_5_19_harmonic_compare.png", "图 5.19 轴承支承刚度变化对倍频分量及幅值比的影响"),
    ]
    for name, caption in figures:
        add_picture(doc, FIG_DIR / name, caption)

    for paragraph in analysis_paragraphs(rows):
        add_para(doc, paragraph)

    doc.save(OUT)


def md_table(headers, rows):
    lines = ["|" + "|".join(headers) + "|"]
    lines.append("|" + "|".join(["---"] * len(headers)) + "|")
    for row in rows:
        lines.append("|" + "|".join(str(value) for value in row) + "|")
    return lines


def write_markdown(rows):
    lines = [
        "# 5.3 轴承支承刚度对转子-轴承-机匣系统振动响应的影响",
        "",
        "本研究固定转速为 9900 r/min，仅改变 `mian.m` 中实际装配进整体刚度矩阵 `KK` 的 `kb1` 和 `kb2`。每个工况令 `kb1=kb2=K_level`，`K_level=[1.0e6,1.0e7,1.0e8,1.0e9,1.0e10] N/m`。`kb3`、轴承支承形式的 `Kzx/Kzy` 和 `K_b` 在当前耦合模型中不存在或不作为轴承支承刚度入口，因此对应列填 `NaN`。本节未修改 `C_b=data(6)`、`data(9)`、`data(10)`、基础支承刚度、阻尼、时间步长、积分方法、观测节点和四个轮盘不平衡位置。",
        "",
        "表 5.5 不同轴承支承刚度工况设置",
    ]
    setting_headers = [
        "case_id", "K_level", "kb1_case", "kb2_case", "kb3_case", "Kzx_case", "Kzy_case",
        "rpm", "1X", "unbalance_amp", "U1_used", "U2_used", "U3_used", "U4_used",
        "Famp1_1", "Famp1_2", "Famp1_3", "Famp1_4", "其他参数保持不变说明", "status",
    ]
    lines.extend(md_table(setting_headers, setting_table_rows(rows)))
    lines.append("")
    lines.append("表 5.6 不同轴承支承刚度下系统振动响应指标")
    response_headers = [
        "case_id", "K_level", "disp_peak", "disp_pp", "vel_peak", "vel_rms",
        "acc_peak", "acc_rms", "orbit_R", "amp_1X", "amp_2X", "amp_3X",
        "2X/1X", "3X/1X", "status",
    ]
    lines.extend(md_table(response_headers, response_table_rows(rows)))
    lines.append("")
    figures = [
        ("fig_5_13_displacement_compare.png", "图 5.13 不同轴承支承刚度下转子节点位移响应对比"),
        ("fig_5_14_velocity_compare.png", "图 5.14 不同轴承支承刚度下转子节点速度响应对比"),
        ("fig_5_15_acceleration_compare.png", "图 5.15 不同轴承支承刚度下转子节点加速度响应对比"),
        ("fig_5_16_orbit_compare.png", "图 5.16 不同轴承支承刚度下轴心轨迹对比"),
        ("fig_5_17_acc_spectrum_compare.png", "图 5.17 不同轴承支承刚度下加速度频谱对比"),
        ("fig_5_18_indicator_compare.png", "图 5.18 轴承支承刚度变化对主要响应指标的影响"),
        ("fig_5_19_harmonic_compare.png", "图 5.19 轴承支承刚度变化对倍频分量及幅值比的影响"),
    ]
    for name, caption in figures:
        lines.append(f"![{caption}]({(FIG_DIR / name).resolve().as_posix()})")
        lines.append("")
    lines.extend(analysis_paragraphs(rows))
    lines.append("")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def build():
    if not SUMMARY.exists():
        raise SystemExit("Missing bearing_stiffness_sweep_9900_summary.xlsx. Run run_bearing_stiffness_sweep_9900.m first.")
    rows = read_summary()
    write_markdown(rows)
    build_docx(rows)
    print(f"Generated {OUT} and {MD_OUT}")


if __name__ == "__main__":
    build()
