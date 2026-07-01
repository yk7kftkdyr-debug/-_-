from pathlib import Path
import csv
import math

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\常用文件\程序\程序\Newmark解故障轴承")
OUT = Path(r"D:\Desktop\第五章_修改后.docx")


def read_csv(name):
    path = ROOT / name
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    return rows


def f(row, key):
    value = row.get(key, "")
    try:
        return float(value)
    except Exception:
        return float("nan")


def sci(value, digits=3):
    if value is None:
        return "NaN"
    try:
        x = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(x):
        return "NaN"
    return f"{x:.{digits}e}"


def num(value, digits=4):
    try:
        x = float(value)
    except Exception:
        return str(value)
    if not math.isfinite(x):
        return "NaN"
    return f"{x:.{digits}g}"


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, size=10.5, bold=False, align=None, before=0, after=5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.28) if not bold else None
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level):
    size = {1: 16, 2: 14, 3: 12.5}.get(level, 12)
    p = add_para(doc, text, size=size, bold=True, before=8 if level > 1 else 0, after=8, color=(31, 78, 121))
    p.paragraph_format.first_line_indent = None
    return p


def cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, caption, headers, rows, font_size=7.2):
    cap = add_para(doc, caption, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    cap.paragraph_format.first_line_indent = None
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=font_size)
        cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side in ("top", "left", "bottom", "right"):
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), "80")
                node.set(qn("w:type"), "dxa")
    add_para(doc, "", after=2)
    return table


def add_picture(doc, rel_path, caption, width=6.35):
    path = ROOT / rel_path
    if not path.exists():
        add_para(doc, f"{caption}（图片文件未找到：{path.name}）", align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = None
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    cap = add_para(doc, caption, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    cap.paragraph_format.first_line_indent = None


def status_text(kind, value):
    if kind == "unbalance":
        if value == 0:
            return "零不平衡基准"
        if value < 5:
            return "健康小振动"
        return "较强不平衡可接受"
    if kind == "load":
        if value <= 2500:
            return "健康/正常载荷"
        if value <= 5000:
            return "中等外载荷"
        return "高载荷接近异常"
    return "合理计算工况"


def build():
    unb = read_csv("unbalance_sweep_9900_summary.csv")
    load = read_csv("external_load_sweep_9900_summary.csv")
    stiff = read_csv("bearing_stiffness_sweep_9900_summary.csv")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    add_heading(doc, "5 关键参数对转子—轴承—机匣系统振动响应的影响分析", 1)
    add_para(doc, "本章基于当前高压转子—轴承—机匣耦合模型，对不平衡量、外加载荷和轴承支承刚度三个关键参数开展单因素分析。为避免概念混淆，本章所有位移、速度、加速度、轴心轨迹和频谱分析对象均统一表述为转子轴承处节点响应，即转子在前后轴承支承位置附近的响应。工程合理性判断采用高压压气机侧轴承座/支撑位置常用经验量级作为参考：速度正常约为 2.0-4.5 mm/s RMS，警戒参考约 7 mm/s，停机参考约 11-12 mm/s；结构振幅峰峰值常见合理范围约 15-50 μm；低频等效加速度约 2-4 m/s^2，健康高频或包络加速度可约 5-30 m/s^2，滚动体损伤冲击时局部峰值可能达到 50-150 m/s^2 以上。")

    add_heading(doc, "5.1 不平衡量对转子—轴承—机匣系统振动响应的影响", 2)
    add_para(doc, "转子不平衡是旋转机械最常见的同步激励源之一，其本质是转子质心与回转中心不重合。偏心质量随转子转动产生周期性离心力，使转子轴承处节点产生以转频 1X 为主的振动响应。实际工程中，不平衡可能来自材料密度不均、加工或装配偏心、长期运行后的沉积与磨损，以及连接件松动或预紧不均等因素。")
    add_para(doc, "本文采用集中质量偏心建模方法描述四个轮盘的不平衡激励。第 i 个轮盘的不平衡量记为 U_i，其离心力幅值为 F_i=U_i*w_i^2，其中 w_i 为转子角速度。由于本节固定转速为 9900 r/min，对应 1X=165 Hz，不同工况下不平衡力幅值的变化仅由 U_i 的比例放大决定。")

    add_heading(doc, "5.1.1 不平衡激励对转子—轴承—机匣系统振动响应的影响", 3)
    add_para(doc, "本节保留原有不平衡量工况 scale=0、0.5、1、2、5。四个轮盘不平衡位置保持为 loca=[4, 6, 8, 14]，各轮盘不平衡量按相同比例系数整体放大或减小。对每组工况提取最后五个转周期的稳态响应，统计转子轴承处节点的位移、速度、加速度、轴心轨迹半径以及 1X、2X、3X 频谱分量。")

    unb_set_rows = []
    for r in unb:
        scale = f(r, "unbalance_scale")
        unb_set_rows.append([
            num(scale), sci(f(r, "U1")), sci(f(r, "U2")), sci(f(r, "U3")), sci(f(r, "U4")),
            sci(f(r, "Famp1_1")), sci(f(r, "Famp1_2")), sci(f(r, "Famp1_3")), sci(f(r, "Famp1_4")),
            status_text("unbalance", scale),
        ])
    add_table(doc, "表 5.1 不同不平衡量工况设置", ["scale", "U1/(kg·m)", "U2/(kg·m)", "U3/(kg·m)", "U4/(kg·m)", "F1/N", "F2/N", "F3/N", "F4/N", "工程属性"], unb_set_rows, 6.5)

    unb_resp_rows = []
    for r in unb:
        scale = f(r, "unbalance_scale")
        unb_resp_rows.append([
            num(scale), f"{f(r,'disp_pp')*1e6:.3f}", f"{f(r,'vel_rms')*1e3:.3f}",
            f"{f(r,'acc_rms'):.3f}", f"{f(r,'orbit_radius_max')*1e6:.3f}",
            f"{f(r,'amp_1X'):.3f}", f"{f(r,'ratio_2X_1X'):.4f}", f"{f(r,'ratio_3X_1X'):.4f}",
            status_text("unbalance", scale),
        ])
    add_table(doc, "表 5.2 不同不平衡量下转子轴承处响应指标", ["scale", "位移p-p/μm", "速度RMS/(mm/s)", "加速度RMS/(m/s^2)", "轨迹半径/μm", "1X幅值", "2X/1X", "3X/1X", "判断"], unb_resp_rows, 7.0)

    add_para(doc, "由表 5.2 和图 5.1-图 5.6 可见，随着不平衡量增大，不平衡力同步增大，转子轴承处节点的位移、速度、加速度和轴心轨迹尺度总体增大。除 scale=0 的零激励基准外，位移峰峰值由约 0.84 μm 增至 12.48 μm，速度 RMS 由约 0.29 mm/s 增至 4.36 mm/s，加速度 RMS 由约 0.31 m/s^2 增至 4.56 m/s^2。该趋势说明当前模型对同步不平衡激励具有明确的幅值敏感性，且响应仍处于小幅振动范围。")
    add_para(doc, "频谱结果表明，各非零不平衡工况下 1X 分量始终占主导，2X/1X 与 3X/1X 幅值比较小，仅随不平衡量增大出现轻微上升。这说明系统主要受同步离心激励控制，轴承非线性接触和支承调制尚未发展为强倍频响应，也未表现出明显碰摩或强非线性失稳特征。")
    add_para(doc, "从工程合理性看，本组工况整体合理。低倍不平衡工况属于健康小振动状态；scale=5 时速度 RMS 约为 4.36 mm/s，接近正常运行上限但仍低于约 7 mm/s 的警戒参考值，位移峰峰值约 12.48 μm，仍低于 15-50 μm 的常见结构振幅范围。因此 scale=5 可作为较强不平衡下的可接受响应分析工况，而不是异常失稳工况。")

    add_picture(doc, "unbalance_sweep_9900_figures/fig_5_1_displacement_compare.png", "图 5.1 不同不平衡量下转子轴承处位移响应对比")
    add_picture(doc, "unbalance_sweep_9900_figures/fig_5_2_velocity_compare.png", "图 5.2 不同不平衡量下转子轴承处速度响应对比")
    add_picture(doc, "unbalance_sweep_9900_figures/fig_5_3_acceleration_compare.png", "图 5.3 不同不平衡量下转子轴承处加速度响应对比")
    add_picture(doc, "unbalance_sweep_9900_figures/fig_5_4_orbit_compare.png", "图 5.4 不同不平衡量下转子轴承处轴心轨迹对比")
    add_picture(doc, "unbalance_sweep_9900_figures/fig_5_5_acc_spectrum_compare.png", "图 5.5 不同不平衡量下转子轴承处加速度频谱对比")
    add_picture(doc, "unbalance_sweep_9900_figures/fig_5_6_indicator_compare.png", "图 5.6 不平衡量变化对转子响应指标的影响")

    add_heading(doc, "5.2 外加载荷对转子—轴承—机匣系统振动响应的影响", 2)
    add_para(doc, "外加载荷是高速转子系统运行过程中的重要外部激励来源。气动力、重力、联轴器载荷、齿轮啮合力以及轴向推力等因素均可能改变转子静平衡位置、轴承受载区和滚动体接触状态，进而影响转子轴承处节点的位移、速度、加速度和频谱结构。")
    add_para(doc, "当前程序中外加载荷通过 F_bearing.m 的 data(10) 进入模型，其形式为 F_r=[data(10);0;data(10);0]，等效作用于前后两个轴承支承位置的 x 向径向载荷。因此，本节 external_load_N 表示两个轴承支承位置各自承受的等效径向载荷，而不是系统总载荷的分配值。")

    add_heading(doc, "5.2.1 外加载荷作用下转子—轴承—机匣系统动力学响应分析", 3)
    add_para(doc, "本节保留外加载荷工况 0 N、1000 N、2500 N、5000 N 和 7500 N，其余参数保持不变。根据工程量级，将 0-2500 N 视为健康或正常载荷工况，5000 N 视为中等外载荷工况，7500 N 视为高外载荷或接近异常工况。")

    load_set_rows = []
    for r in load:
        value = f(r, "external_load_N")
        load_set_rows.append([r["case_id"], num(value), "轴承1、轴承2", "+x 向", "9900", "165", status_text("load", value)])
    add_table(doc, "表 5.3 不同外加载荷工况设置", ["工况", "外加载荷/N", "作用位置", "方向", "转速/(r/min)", "1X/Hz", "工程属性"], load_set_rows, 8.0)

    load_resp_rows = []
    for r in load:
        value = f(r, "external_load_N")
        load_resp_rows.append([
            r["case_id"], num(value), f"{f(r,'disp_pp')*1e6:.2f}", f"{f(r,'vel_rms')*1e3:.2f}",
            f"{f(r,'acc_rms'):.2f}", f"{f(r,'orbit_radius_max')*1e6:.2f}",
            f"{f(r,'amp_1X'):.3f}", f"{f(r,'ratio_2X_1X'):.4f}", f"{f(r,'ratio_3X_1X'):.4f}",
            status_text("load", value),
        ])
    add_table(doc, "表 5.4 不同外加载荷下转子轴承处响应指标", ["工况", "载荷/N", "位移p-p/μm", "速度RMS/(mm/s)", "加速度RMS/(m/s^2)", "轨迹半径/μm", "1X幅值", "2X/1X", "3X/1X", "判断"], load_resp_rows, 7.0)

    add_para(doc, "由表 5.4 和图 5.7-图 5.13 可见，外加载荷增大后，转子轴承处节点响应总体增强。位移峰峰值由约 1.74 μm 增至 47.13 μm，轴心轨迹最大半径由约 0.87 μm 增至 24.22 μm，整体随载荷增大而扩大。速度 RMS 由约 0.62 mm/s 增至 11.64 mm/s，加速度 RMS 由约 0.65 m/s^2 增至 9.41 m/s^2，但中间工况存在非单调变化，说明外载荷不仅改变响应幅值，也改变轴承接触状态和倍频调制关系。")
    add_para(doc, "频域结果显示，1X 同步分量仍是主要成分，但 2X/1X 和 3X/1X 幅值比随载荷增大总体升高。尤其在 5000 N 和 7500 N 工况下，倍频占比明显高于低载荷工况，表明轴承支承接触非线性和载荷调制效应增强。该现象符合滚动轴承在较高径向载荷下接触区变化、有效游隙减小和支承反力重新分配的动力学特征。")
    add_para(doc, "从工程合理性看，0-2500 N 工况的速度 RMS 处于健康或正常范围内，可作为稳定运行状态分析；5000 N 工况虽然速度未超过警戒线，但位移和倍频分量已明显增强，适合用于分析中等外载荷下的非线性增强。7500 N 工况下速度 RMS 达到约 11.64 mm/s，已接近 11-12 mm/s 的停机参考范围，位移峰峰值约 47 μm，也接近结构振幅合理范围上限。因此，7500 N 不宜定义为健康运行状态，更适合作为高外载荷下的重载非线性响应或接近异常工况分析。该工况可保留用于说明外载荷过大时系统响应显著增强，但不能作为正常工况外推。")

    add_picture(doc, "external_load_sweep_9900_report_figures/fig_5_7_displacement_compare.png", "图 5.7 不同外加载荷下转子轴承处位移响应对比")
    add_picture(doc, "external_load_sweep_9900_report_figures/fig_5_8_velocity_compare.png", "图 5.8 不同外加载荷下转子轴承处速度响应对比")
    add_picture(doc, "external_load_sweep_9900_report_figures/fig_5_9_acceleration_compare.png", "图 5.9 不同外加载荷下转子轴承处加速度响应对比")
    add_picture(doc, "external_load_sweep_9900_report_figures/fig_5_10_orbit_compare.png", "图 5.10 不同外加载荷下转子轴承处轴心轨迹对比")
    add_picture(doc, "external_load_sweep_9900_report_figures/fig_5_11_acc_spectrum_compare.png", "图 5.11 不同外加载荷下转子轴承处加速度频谱对比")
    add_picture(doc, "external_load_sweep_9900_report_figures/fig_5_13_harmonic_compare.png", "图 5.12 外加载荷变化对倍频分量及幅值比的影响")
    add_picture(doc, "external_load_sweep_9900_report_figures/fig_5_12_indicator_compare.png", "图 5.13 外加载荷变化对转子响应指标的影响")

    add_heading(doc, "5.3 轴承支承刚度对转子—轴承—机匣系统振动响应的影响", 2)
    add_para(doc, "轴承支承刚度直接影响高速转子系统的支撑特性、临界转速、模态分布和不平衡响应。支承刚度过低时，转子支撑柔性增强，轴心轨迹和位移响应可能显著扩大；支承刚度过高时，系统约束增强，位移通常减小，但局部接触状态和高频响应也可能发生变化。因此，轴承支承刚度分析必须首先排除明显非物理工况，再讨论稳定工程范围内的规律。")
    add_para(doc, "前期试算中采用的 1e6 N/m 和 1e7 N/m 工况导致位移达到 0.25 m 量级、速度达到 100 m/s 量级、加速度达到 10^6 m/s^2 量级，说明支承约束严重不足，已不符合高压转子轴承支承刚度实际范围。因此，这两个低刚度工况不再作为正式结果参与趋势分析，仅作为前期试算中被排除的异常工况。")

    add_heading(doc, "5.3.1 轴承支承刚度变化对转子—轴承—机匣系统振动特性的影响", 3)
    add_para(doc, "根据程序稳定性和高压转子工程可解释范围，本次重新设置轴承支承刚度为 5e7、1e8、2e8、5e8、1e9 和 2e9 N/m。该范围集中在 1e8-1e9 N/m 附近，同时覆盖较软支承和较硬支承两侧，用于分析支承刚度变化对转子轴承处节点响应的影响。所有工况均完成全时长积分，未出现 10^-1 m 量级位移或 10^6 m/s^2 量级加速度。")

    stiff_set_rows = []
    for r in stiff:
        stiff_set_rows.append([r["case_id"], sci(f(r, "K_level"), 2), sci(f(r, "kb1_case"), 2), sci(f(r, "kb2_case"), 2), "9900", "165", r["simulation_status"], "正式合理工况"])
    add_table(doc, "表 5.5 不同轴承支承刚度工况设置", ["工况", "K_level/(N/m)", "kb1/(N/m)", "kb2/(N/m)", "转速/(r/min)", "1X/Hz", "状态", "说明"], stiff_set_rows, 7.3)

    stiff_resp_rows = []
    for r in stiff:
        stiff_resp_rows.append([
            r["case_id"], sci(f(r, "K_level"), 2), f"{f(r,'disp_pp')*1e6:.3f}", f"{f(r,'vel_rms')*1e3:.3f}",
            f"{f(r,'acc_rms'):.3f}", f"{f(r,'orbit_radius_max')*1e6:.3f}", f"{f(r,'amp_1X'):.3f}",
            f"{f(r,'ratio_2X_1X'):.4f}", f"{f(r,'ratio_3X_1X'):.4f}", r["simulation_status"],
        ])
    add_table(doc, "表 5.6 不同轴承支承刚度下转子轴承处响应指标", ["工况", "K/(N/m)", "位移p-p/μm", "速度RMS/(mm/s)", "加速度RMS/(m/s^2)", "轨迹半径/μm", "1X幅值", "2X/1X", "3X/1X", "状态"], stiff_resp_rows, 7.0)

    add_para(doc, "由表 5.6 和图 5.14-图 5.20 可见，重新选取的刚度范围内响应均处于高速转子工程可解释量级。K=5e7-2e8 N/m 时，位移峰峰值约为 1.18-3.42 μm，速度 RMS 约为 0.32-1.20 mm/s，加速度 RMS 约为 0.32-1.23 m/s^2；当刚度继续增大至 5e8-2e9 N/m 时，位移、速度、加速度和轴心轨迹半径总体降低。")
    add_para(doc, "需要注意的是，K=2e8 N/m 工况响应略高于 1e8 N/m 工况，说明在当前转速和非线性轴承接触模型下，刚度变化并非严格线性单调。该局部峰值可能与支承刚度改变后系统模态位置、接触反力分布和 1X 同步响应耦合有关。进入 5e8 N/m 及以上较高刚度区间后，支承约束增强，响应幅值明显降低，表明合理提高轴承支承刚度能够有效抑制转子轴承处横向振动。")
    add_para(doc, "频域结果显示，各正式刚度工况仍以 1X 同步响应为主，2X/1X 和 3X/1X 幅值比较低，没有出现强倍频或高频冲击主导现象。从工程合理性看，本组新刚度工况全部低于速度警戒参考值，位移峰峰值也远低于 15-50 μm 的常见结构振幅范围，属于可用于论文或报告趋势分析的合理计算工况。原 1e6 和 1e7 N/m 工况仅可作为参数敏感性或异常边界说明，不作为稳定工程工况进行后续规律分析。")

    add_picture(doc, "bearing_stiffness_sweep_9900_figures/fig_5_13_displacement_compare.png", "图 5.14 不同轴承支承刚度下转子轴承处位移响应对比")
    add_picture(doc, "bearing_stiffness_sweep_9900_figures/fig_5_14_velocity_compare.png", "图 5.15 不同轴承支承刚度下转子轴承处速度响应对比")
    add_picture(doc, "bearing_stiffness_sweep_9900_figures/fig_5_15_acceleration_compare.png", "图 5.16 不同轴承支承刚度下转子轴承处加速度响应对比")
    add_picture(doc, "bearing_stiffness_sweep_9900_figures/fig_5_16_orbit_compare.png", "图 5.17 不同轴承支承刚度下转子轴承处轴心轨迹对比")
    add_picture(doc, "bearing_stiffness_sweep_9900_figures/fig_5_17_acc_spectrum_compare.png", "图 5.18 不同轴承支承刚度下转子轴承处加速度频谱对比")
    add_picture(doc, "bearing_stiffness_sweep_9900_figures/fig_5_19_harmonic_compare.png", "图 5.19 轴承支承刚度变化对倍频分量及幅值比的影响")
    add_picture(doc, "bearing_stiffness_sweep_9900_figures/fig_5_18_indicator_compare.png", "图 5.20 轴承支承刚度变化对转子响应指标的影响")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
