from copy import deepcopy
from pathlib import Path
import os
import re
import shutil
import tempfile
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"D:\Desktop\高压转子—轴承—机匣系统关键参数对振动响应的影响分析.docx")
TPL = Path(r"D:\apps\xwechat_files\wxid_u5f8107b9hgc22_daec\msg\file\2026-05\1.船用混合动力传动系统动力学仿真及测试技术调研报告(3)(1).docx")
OUT = Path.cwd() / "高压转子系统报告_模板版.docx"
REPORT_NAME = "高压转子—轴承—机匣系统关键参数对振动响应的影响分析"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}
for prefix, uri in NS.items():
    ET.register_namespace(prefix if prefix not in ("ct", "rel") else "", uri)


def qn_xml(prefix, tag):
    return f"{{{NS[prefix]}}}{tag}"


def read_xml(zip_path, name):
    with zipfile.ZipFile(zip_path) as z:
        return ET.fromstring(z.read(name))


def write_xml(path, root):
    ET.ElementTree(root).write(path, encoding="utf-8", xml_declaration=True)


def get_style_ids(template_docx):
    doc = Document(str(template_docx))
    ids = {"Normal": "Normal", "Heading 1": "Heading1", "Heading 2": "Heading2", "Heading 3": "Heading3", "图": "Caption"}
    for style in doc.styles:
        if style.name in ids:
            ids[style.name] = style.style_id
    return ids


def paragraph_text(p):
    return "".join(t.text or "" for t in p.findall(".//w:t", NS)).strip()


def ensure_pstyle(p, style_id):
    ppr = p.find("w:pPr", NS)
    if ppr is None:
        ppr = ET.Element(qn_xml("w", "pPr"))
        p.insert(0, ppr)
    pstyle = ppr.find("w:pStyle", NS)
    if pstyle is None:
        pstyle = ET.Element(qn_xml("w", "pStyle"))
        ppr.insert(0, pstyle)
    pstyle.set(qn_xml("w", "val"), style_id)


def set_para_center(p):
    ppr = p.find("w:pPr", NS)
    if ppr is None:
        ppr = ET.Element(qn_xml("w", "pPr"))
        p.insert(0, ppr)
    jc = ppr.find("w:jc", NS)
    if jc is None:
        jc = ET.Element(qn_xml("w", "jc"))
        ppr.append(jc)
    jc.set(qn_xml("w", "val"), "center")


def classify_and_style_body_elements(elements, style_ids):
    h1 = style_ids.get("Heading 1", "Heading1")
    h2 = style_ids.get("Heading 2", "Heading2")
    h3 = style_ids.get("Heading 3", "Heading3")
    normal = style_ids.get("Normal", "Normal")
    caption = style_ids.get("图", "Caption")

    for el in elements:
        if el.tag != qn_xml("w", "p"):
            continue
        txt = paragraph_text(el)
        if not txt:
            continue
        if re.match(r"^第\s*\d+\s*章", txt):
            ensure_pstyle(el, h1)
        elif re.match(r"^\d+\.\d+\.\d+(\s|$)", txt):
            ensure_pstyle(el, h3)
        elif re.match(r"^\d+\.\d+(\s|$)", txt):
            ensure_pstyle(el, h2)
        elif re.match(r"^(图|表)\s*\d+", txt):
            ensure_pstyle(el, caption)
            set_para_center(el)
        elif el.find(".//m:oMath", {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}) is not None:
            set_para_center(el)
        else:
            ppr = el.find("w:pPr", NS)
            pstyle = ppr.find("w:pStyle", NS) if ppr is not None else None
            if pstyle is not None and pstyle.get(qn_xml("w", "val")) in ("Normal", "NormalWeb", "a3"):
                pstyle.set(qn_xml("w", "val"), normal)


def make_toc_elements(style_ids):
    normal = style_ids.get("Normal", "Normal")
    toc_title = style_ids.get("TOC 标题1", normal)

    note = ET.Element(qn_xml("w", "p"))
    ensure_pstyle(note, normal)
    nr = ET.SubElement(note, qn_xml("w", "r"))
    nt = ET.SubElement(nr, qn_xml("w", "t"))
    nt.text = "提示：打开 Word 后请右键目录并选择更新域。"

    title = ET.Element(qn_xml("w", "p"))
    ensure_pstyle(title, toc_title)
    set_para_center(title)
    tr = ET.SubElement(title, qn_xml("w", "r"))
    tt = ET.SubElement(tr, qn_xml("w", "t"))
    tt.text = "目  录"

    # Complex field: TOC \o "1-3" \h \z \u
    p = ET.Element(qn_xml("w", "p"))
    r1 = ET.SubElement(p, qn_xml("w", "r"))
    fld_begin = ET.SubElement(r1, qn_xml("w", "fldChar"))
    fld_begin.set(qn_xml("w", "fldCharType"), "begin")
    r2 = ET.SubElement(p, qn_xml("w", "r"))
    instr = ET.SubElement(r2, qn_xml("w", "instrText"))
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r3 = ET.SubElement(p, qn_xml("w", "r"))
    fld_sep = ET.SubElement(r3, qn_xml("w", "fldChar"))
    fld_sep.set(qn_xml("w", "fldCharType"), "separate")
    r4 = ET.SubElement(p, qn_xml("w", "r"))
    placeholder = ET.SubElement(r4, qn_xml("w", "t"))
    placeholder.text = "目录域将在 Word 中更新。"
    r5 = ET.SubElement(p, qn_xml("w", "r"))
    fld_end = ET.SubElement(r5, qn_xml("w", "fldChar"))
    fld_end.set(qn_xml("w", "fldCharType"), "end")

    page_break = ET.Element(qn_xml("w", "p"))
    br_run = ET.SubElement(page_break, qn_xml("w", "r"))
    br = ET.SubElement(br_run, qn_xml("w", "br"))
    br.set(qn_xml("w", "type"), "page")
    return [note, title, p, page_break]


def max_rid(rels_root):
    max_id = 0
    for rel in rels_root.findall("rel:Relationship", NS):
        rid = rel.get("Id", "")
        m = re.match(r"rId(\d+)$", rid)
        if m:
            max_id = max(max_id, int(m.group(1)))
    return max_id


def copy_related_parts(src_zip, work_dir, out_rels_root, src_body_elements):
    src_rels = read_xml(SRC, "word/_rels/document.xml.rels")
    existing_targets = {rel.get("Target") for rel in out_rels_root.findall("rel:Relationship", NS)}
    next_id = max_rid(out_rels_root) + 1
    rid_map = {}

    used_rids = set()
    for el in src_body_elements:
        for node in el.iter():
            for attr, val in node.attrib.items():
                if attr.startswith(f"{{{NS['r']}}}"):
                    used_rids.add(val)

    for rel in src_rels.findall("rel:Relationship", NS):
        old_id = rel.get("Id")
        if old_id not in used_rids:
            continue
        new_id = f"rId{next_id}"
        next_id += 1
        rid_map[old_id] = new_id
        new_rel = deepcopy(rel)
        new_rel.set("Id", new_id)
        target = rel.get("Target")
        mode = rel.get("TargetMode")
        if mode != "External" and target:
            src_part = "word/" + target if not target.startswith("/") else target.lstrip("/")
            src_part = os.path.normpath(src_part).replace("\\", "/")
            basename = Path(target).name
            subdir = str(Path(target).parent).replace("\\", "/")
            if subdir == ".":
                subdir = ""
            new_target = target
            if target in existing_targets or (work_dir / "word" / target).exists():
                new_target = f"{subdir}/src_{old_id}_{basename}" if subdir else f"src_{old_id}_{basename}"
                new_target = new_target.replace("\\", "/")
                new_rel.set("Target", new_target)
            dst = work_dir / "word" / new_target
            dst.parent.mkdir(parents=True, exist_ok=True)
            with src_zip.open(src_part) as fsrc:
                dst.write_bytes(fsrc.read())
            existing_targets.add(new_target)
        out_rels_root.append(new_rel)
    return rid_map


def remap_rids(elements, rid_map):
    if not rid_map:
        return
    for el in elements:
        for node in el.iter():
            for attr, val in list(node.attrib.items()):
                if attr.startswith(f"{{{NS['r']}}}") and val in rid_map:
                    node.set(attr, rid_map[val])


def merge_content_types(work_dir):
    out_ct = ET.parse(work_dir / "[Content_Types].xml").getroot()
    src_ct = read_xml(SRC, "[Content_Types].xml")
    existing_defaults = {(d.get("Extension"), d.get("ContentType")) for d in out_ct.findall("ct:Default", NS)}
    existing_exts = {d.get("Extension") for d in out_ct.findall("ct:Default", NS)}
    for d in src_ct.findall("ct:Default", NS):
        ext = d.get("Extension")
        if ext not in existing_exts:
            out_ct.append(deepcopy(d))
            existing_exts.add(ext)
            existing_defaults.add((ext, d.get("ContentType")))
    # Internal copied images normally use defaults. Copy source overrides only if
    # their part names already exist after copying; this avoids stale overrides.
    existing_parts = {"/" + str(p.relative_to(work_dir)).replace("\\", "/") for p in work_dir.rglob("*") if p.is_file()}
    existing_overrides = {o.get("PartName") for o in out_ct.findall("ct:Override", NS)}
    for o in src_ct.findall("ct:Override", NS):
        part = o.get("PartName")
        if part in existing_parts and part not in existing_overrides:
            out_ct.append(deepcopy(o))
            existing_overrides.add(part)
    write_xml(work_dir / "[Content_Types].xml", out_ct)


def set_update_fields(work_dir):
    settings_path = work_dir / "word" / "settings.xml"
    if not settings_path.exists():
        return
    root = ET.parse(settings_path).getroot()
    update = root.find("w:updateFields", NS)
    if update is None:
        update = ET.Element(qn_xml("w", "updateFields"))
        root.append(update)
    update.set(qn_xml("w", "val"), "true")
    write_xml(settings_path, root)


def replace_header_footer_report_name(work_dir):
    candidates = list((work_dir / "word").glob("header*.xml")) + list((work_dir / "word").glob("footer*.xml"))
    replacements = [
        "船用混合动力传动系统动力学仿真及测试技术调研报告",
        "船用混合动力传动系统动力学仿真及测试技术",
        "混合动力传动系统动力学仿真及测试技术调研报告",
    ]
    for path in candidates:
        text = path.read_text(encoding="utf-8", errors="ignore")
        new_text = text
        for old in replacements:
            new_text = new_text.replace(old, REPORT_NAME)
        if new_text != text:
            path.write_text(new_text, encoding="utf-8")


def build_package():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if not TPL.exists():
        raise FileNotFoundError(TPL)
    style_ids = get_style_ids(TPL)
    with tempfile.TemporaryDirectory() as td:
        work = Path(td)
        with zipfile.ZipFile(TPL) as z:
            z.extractall(work)
        template_doc = ET.parse(work / "word" / "document.xml").getroot()
        template_body = template_doc.find("w:body", NS)
        template_sect = template_body.find("w:sectPr", NS)
        if template_sect is None:
            template_sect = ET.Element(qn_xml("w", "sectPr"))

        src_doc = read_xml(SRC, "word/document.xml")
        src_body = src_doc.find("w:body", NS)
        src_elements = [deepcopy(el) for el in list(src_body) if el.tag != qn_xml("w", "sectPr")]
        classify_and_style_body_elements(src_elements, style_ids)

        out_rels_path = work / "word" / "_rels" / "document.xml.rels"
        out_rels_root = ET.parse(out_rels_path).getroot()
        with zipfile.ZipFile(SRC) as src_zip:
            rid_map = copy_related_parts(src_zip, work, out_rels_root, src_elements)
        remap_rids(src_elements, rid_map)
        write_xml(out_rels_path, out_rels_root)

        for child in list(template_body):
            template_body.remove(child)
        for el in make_toc_elements(style_ids):
            template_body.append(el)
        for el in src_elements:
            template_body.append(el)
        template_body.append(deepcopy(template_sect))
        write_xml(work / "word" / "document.xml", template_doc)

        merge_content_types(work)
        set_update_fields(work)
        replace_header_footer_report_name(work)

        if OUT.exists():
            OUT.unlink()
        with zipfile.ZipFile(OUT, "w", compression=zipfile.ZIP_DEFLATED) as out_zip:
            for file in work.rglob("*"):
                if file.is_file():
                    out_zip.write(file, file.relative_to(work).as_posix())


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_border(cell, edge, size=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    element = borders.find(qn("w:" + edge))
    if element is None:
        element = OxmlElement("w:" + edge)
        borders.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:color"), "000000")


def polish_docx():
    doc = Document(str(OUT))
    # Center pictures and equations; style captions; keep text unchanged.
    for p in doc.paragraphs:
        text = p.text.strip()
        has_drawing = bool(p._p.xpath(".//*[local-name()='drawing']"))
        has_math = bool(p._p.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']"))
        if has_drawing or has_math:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if re.match(r"^(图|表)\s*\d+", text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.style = "图"
            except Exception:
                pass

    # Apply three-line table style without changing cell contents.
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                clear_cell_borders(cell)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        if run.font.size is None:
                            run.font.size = Pt(9)
        if table.rows:
            for cell in table.rows[0].cells:
                set_cell_border(cell, "top", 12)
                set_cell_border(cell, "bottom", 8)
            for cell in table.rows[-1].cells:
                set_cell_border(cell, "bottom", 12)

    doc.save(str(OUT))


def inspect_doc(path):
    doc = Document(str(path))
    chapter_text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "sections": len(doc.sections),
        "has_toc_field": "TOC" in zipfile.ZipFile(path).read("word/document.xml").decode("utf-8", errors="ignore"),
        "chapters": [f"第{i}章" in chapter_text for i in range(1, 7)],
        "has_template_body": "混合动力齿轮传动系统动力学建模研究现状" in chapter_text or "油电混合动力推进示意图" in chapter_text,
        "omml_count": zipfile.ZipFile(path).read("word/document.xml").decode("utf-8", errors="ignore").count("<m:oMath"),
    }


def main():
    before_src = (SRC.stat().st_size, SRC.stat().st_mtime_ns)
    before_tpl = (TPL.stat().st_size, TPL.stat().st_mtime_ns)
    build_package()
    polish_docx()
    after_src = (SRC.stat().st_size, SRC.stat().st_mtime_ns)
    after_tpl = (TPL.stat().st_size, TPL.stat().st_mtime_ns)
    src_info = inspect_doc(SRC)
    out_info = inspect_doc(OUT)
    print("source_unchanged", before_src == after_src)
    print("template_unchanged", before_tpl == after_tpl)
    print("source", src_info)
    print("output", out_info)
    print("output_path", OUT)


if __name__ == "__main__":
    main()
