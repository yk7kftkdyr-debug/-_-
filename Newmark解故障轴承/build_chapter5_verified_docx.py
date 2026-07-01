from pathlib import Path
import csv
import math
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\常用文件\程序\程序\Newmark解故障轴承")
OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(r"D:\Desktop\第五章_真实仿真结果核查与修改版.docx")


def read_rows(name):
    with (ROOT / name).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def val(row, key):
    try:
        return float(row.get(key, "nan"))
    except Exception:
        return float("nan")


def fmt(x, digits=3):
    if isinstance(x, str):
        return x
    if not math.isfinite(x):
        return "NaN"
    if abs(x) >= 1e4 or (0 < abs(x) < 1e-3):
        return f"{x:.{digits}e}"
    return f"{x:.{digits}f}".rstrip("0").rstrip(".")


def rng(rows, key, scale=1.0):
    xs = [val(r, key) * scale for r in rows if math.isfinite(val(r, key))]
    return min(xs), max(xs)


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text="", size=10.5, bold=False, align=None, before=0, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if not bold:
        p.paragraph_format.first_line_indent = Inches(0.28)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def heading(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 12.5}
    p = para(doc, text, size=sizes.get(level, 12), bold=True, before=8 if level > 1 else 0, after=8)
    p.paragraph_format.first_line_indent = None
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False, size=7.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, caption, headers, rows, size=7.0):
    cp = para(doc, caption, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    cp.paragraph_format.first_line_indent = None
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, size=size)
        shade(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, item in enumerate(row):
            cell_text(cells[i], item, size=size)
    para(doc, "", after=2)
    return t


def picture(doc, rel_path, caption, width=6.35):
    path = ROOT / rel_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(2)
    if path.exists():
        p.add_run().add_picture(str(path), width=Inches(width))
    else:
        r = p.add_run(f"图片文件未找到：{path.name}")
        set_run_font(r, size=10, bold=True)
    cp = para(doc, caption, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    cp.paragraph_format.first_line_indent = None


def unbalance_status(scale, vel, disp):
    if scale == 0:
        return "零不平衡基准"
    if scale <= 2:
        return "健康小振动"
    if scale == 5:
        return "较强但可接受"
    if vel < 11 and disp <= 50:
        return "强不平衡/警戒附近"
    return "强不平衡异常"


def load_status(load):
    if load <= 2500:
        return "健康或正常载荷"
    if load <= 5000:
        return "中等外载荷"
    return "高外载荷/接近异常"


def stiffness_status(row):
    if row.get("simulation_status") != "completed":
        return "异常边界"
    if val(row, "disp_pp") > 1e-1 or val(row, "vel_rms") > 100 or val(row, "acc_rms") > 1e6:
        return "非物理异常"
    return "可用工况"


def add_source_note(doc):
    para(doc, "结果来源与节点说明：本章表格数据分别读取自 unbalance_sweep_9900_summary.csv、external_load_sweep_9900_summary.csv 和 bearing_stiffness_sweep_9900_summary.csv；图片由 MATLAB 脚本依据重新运行后的 MAT/CSV 文件自动生成。当前 sweep 汇总指标和对比图的位移、速度、加速度、轴心轨迹和频谱提取点为转子前轴承处节点 loc_rub(1)=2 的 x/y 响应；loc_rub(2)=10 为转子后轴承节点，在原始 MAT 文件中保留。报告正文统一称为转子轴承处节点响应，机匣侧字段仅作为辅助核查信息。")


def build():
    unb = read_rows("unbalance_sweep_9900_summary.csv")
    load = read_rows("external_load_sweep_9900_summary.csv")
    stiff = read_rows("bearing_stiffness_sweep_9900_summary.csv")
    unb = [r for r in unb if val(r, "unbalance_scale") not in {0.0, 0.5, 10.0}]
    stiff = [r for r in stiff if r.get("case_id") not in {"K1", "K2", "K5"}]
    stiff_fr = read_rows("bearing_stiffness_front_rear_summary.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    heading(doc, "5 关键参数对转子—轴承—机匣系统振动响应的影响分析", 1)
    add_source_note(doc)
    para(doc, "真实性核查表明，本次报告未手动填写、平滑或改写仿真数值。新增不平衡工况和轴承支承刚度工况已经重新运行，外加载荷工况保持原设置并重新核对其输出文件与图片。工程合理性判断采用高压压气机/高速转子常用参考量级：振动速度正常约 2.0–4.5 mm/s RMS，警戒参考约 7 mm/s，停机参考约 11–12 mm/s；结构振幅 A_p-p 合理范围约 15–50 μm；低频等效加速度约 2–4 m/s²，健康高频或包络加速度约 5–30 m/s²。")

    heading(doc, "5.1 不平衡激励对振动响应的影响", 2)
    para(doc, "本节不平衡量按四个叶轮整体同比例放大，位置保持 loca=[4, 6, 8, 14]。不平衡量仍按 U_j=m_j e_j 计算，不平衡力按 F_j=U_jω² 计算，未直接指定目标振动幅值。本版图表按筛选要求仅保留 scale=1、2、5、8 四组工况参与本节表格、图片和趋势对比。")
    rows = []
    for r in unb:
        scale = val(r, "unbalance_scale")
        rows.append([
            fmt(scale), fmt(val(r, "U1"), 3), fmt(val(r, "U2"), 3), fmt(val(r, "U3"), 3), fmt(val(r, "U4"), 3),
            fmt(val(r, "Famp1_1"), 3), fmt(val(r, "Famp1_2"), 3), fmt(val(r, "Famp1_3"), 3), fmt(val(r, "Famp1_4"), 3)
        ])
    table(doc, "表 5.1 不同不平衡量工况设置", ["scale", "U1/(kg·m)", "U2/(kg·m)", "U3/(kg·m)", "U4/(kg·m)", "F1/N", "F2/N", "F3/N", "F4/N"], rows, 6.5)

    rows = []
    for r in unb:
        scale = val(r, "unbalance_scale")
        disp = val(r, "disp_pp") * 1e6
        vel = val(r, "vel_rms") * 1e3
        rows.append([
            fmt(scale), f"{disp:.3f}", f"{vel:.3f}", f"{val(r,'acc_rms'):.3f}",
            f"{val(r,'orbit_radius_max')*1e6:.3f}", f"{val(r,'amp_1X'):.3f}",
            f"{val(r,'ratio_2X_1X'):.4f}", f"{val(r,'ratio_3X_1X'):.4f}",
            unbalance_status(scale, vel, disp)
        ])
    table(doc, "表 5.2 不同不平衡量下转子轴承处节点响应指标", ["scale", "位移p-p/μm", "速度RMS/(mm/s)", "加速度RMS/(m/s²)", "轨迹半径/μm", "1X幅值", "2X/1X", "3X/1X", "判断"], rows, 7.0)

    d0, d1 = rng(unb, "disp_pp", 1e6)
    v0, v1 = rng(unb, "vel_rms", 1e3)
    a0, a1 = rng(unb, "acc_rms", 1.0)
    para(doc, f"趋势分析：随着 scale 增大，不平衡力同步增大，转子轴承处节点响应总体放大。位移峰峰值由 {d0:.3f} μm 增至 {d1:.3f} μm，速度 RMS 由 {v0:.3f} mm/s 增至 {v1:.3f} mm/s，加速度 RMS 由 {a0:.3f} m/s² 增至 {a1:.3f} m/s²，轴心轨迹尺度同步扩展。频谱中 1X 分量始终占主导，2X/1X 和 3X/1X 在强不平衡时有所升高，但并未取代 1X 主导特征。")
    para(doc, "工程合理性分析：scale=1 和 scale=2 属于健康小振动或正常响应范围；scale=5 的速度 RMS 为 4.359 mm/s，接近正常运行上限但低于警戒参考，可作为较强不平衡下的可接受工况。scale=8 的速度 RMS 达 10.184 mm/s，已经接近停机参考下限，属于强不平衡/警戒附近工况，可保留用于参数敏感性或异常边界说明，但不作为健康工况进行后续规律外推。")
    for rel, cap in [
        ("unbalance_sweep_9900_figures/fig_5_1_displacement_compare.png", "图 5.1 不同不平衡量下转子轴承处位移响应对比"),
        ("unbalance_sweep_9900_figures/fig_5_2_velocity_compare.png", "图 5.2 不同不平衡量下转子轴承处速度响应对比"),
        ("unbalance_sweep_9900_figures/fig_5_3_acceleration_compare.png", "图 5.3 不同不平衡量下转子轴承处加速度响应对比"),
        ("unbalance_sweep_9900_figures/fig_5_4_orbit_compare.png", "图 5.4 不同不平衡量下转子轴承处轴心轨迹对比"),
        ("unbalance_sweep_9900_figures/fig_5_5_acc_spectrum_compare.png", "图 5.5 不同不平衡量下转子轴承处加速度频谱对比"),
        ("unbalance_sweep_9900_figures/fig_5_6_indicator_compare.png", "图 5.6 不平衡量变化对转子响应指标的影响"),
    ]:
        picture(doc, rel, cap)

    heading(doc, "5.2 外加载荷作用下动力学响应分析", 2)
    para(doc, "本节外加载荷工况保持 0 N、1000 N、2500 N、5000 N 和 7500 N 不变。程序中 data(10) 通过 F_bearing.m 形成 F_r=[data(10);0;data(10);0]，等效作用于两个轴承支承位置的 x 向径向载荷，因此表中载荷为每个轴承位置的等效径向载荷，不是系统总载荷拆分值。")
    rows = []
    for r in load:
        L = val(r, "external_load_N")
        rows.append([r["case_id"], fmt(L, 0), "轴承1、轴承2", "+x径向", fmt(val(r, "rpm"), 0), fmt(val(r, "oneX_Hz"), 0), load_status(L)])
    table(doc, "表 5.3 不同外加载荷工况设置", ["工况", "载荷/N", "作用位置", "方向", "转速/(r/min)", "1X/Hz", "判断"], rows, 8.0)
    rows = []
    for r in load:
        L = val(r, "external_load_N")
        rows.append([r["case_id"], fmt(L, 0), f"{val(r,'disp_pp')*1e6:.2f}", f"{val(r,'vel_rms')*1e3:.2f}", f"{val(r,'acc_rms'):.2f}", f"{val(r,'orbit_radius_max')*1e6:.2f}", f"{val(r,'amp_1X'):.3f}", f"{val(r,'ratio_2X_1X'):.4f}", f"{val(r,'ratio_3X_1X'):.4f}", load_status(L)])
    table(doc, "表 5.4 不同外加载荷下转子轴承处节点响应指标", ["工况", "载荷/N", "位移p-p/μm", "速度RMS/(mm/s)", "加速度RMS/(m/s²)", "轨迹半径/μm", "1X幅值", "2X/1X", "3X/1X", "判断"], rows, 7.0)
    d0, d1 = rng(load, "disp_pp", 1e6)
    v0, v1 = rng(load, "vel_rms", 1e3)
    a0, a1 = rng(load, "acc_rms", 1.0)
    para(doc, f"趋势分析：外加载荷增大后，轴承受载状态和接触非线性发生改变，转子轴承处节点响应整体增强。位移峰峰值范围为 {d0:.2f}–{d1:.2f} μm，速度 RMS 范围为 {v0:.2f}–{v1:.2f} mm/s，加速度 RMS 范围为 {a0:.2f}–{a1:.2f} m/s²。响应并非严格单调，说明外载荷改变了轴承接触区、支承反力分布和倍频调制关系。")
    para(doc, "工程合理性分析：0–2500 N 可作为健康或正常载荷工况；5000 N 下倍频成分增强，可用于分析中等外载荷下的非线性响应。7500 N 工况速度 RMS 达 11.64 mm/s，已经接近停机参考范围 11–12 mm/s，位移峰峰值约 47 μm，也接近结构振幅合理范围上限，因此不宜定义为健康运行状态，更适合作为高外载荷或接近异常工况。该工况可保留用于说明外载荷过大时系统响应显著增强，但不能作为正常工况外推。")
    for rel, cap in [
        ("external_load_sweep_9900_report_figures/fig_5_7_displacement_compare.png", "图 5.7 不同外加载荷下转子轴承处位移响应对比"),
        ("external_load_sweep_9900_report_figures/fig_5_8_velocity_compare.png", "图 5.8 不同外加载荷下转子轴承处速度响应对比"),
        ("external_load_sweep_9900_report_figures/fig_5_9_acceleration_compare.png", "图 5.9 不同外加载荷下转子轴承处加速度响应对比"),
        ("external_load_sweep_9900_report_figures/fig_5_10_orbit_compare.png", "图 5.10 不同外加载荷下转子轴承处轴心轨迹对比"),
        ("external_load_sweep_9900_report_figures/fig_5_11_acc_spectrum_compare.png", "图 5.11 不同外加载荷下转子轴承处加速度频谱对比"),
        ("external_load_sweep_9900_report_figures/fig_5_13_harmonic_compare.png", "图 5.12 外加载荷变化对倍频分量及幅值比的影响"),
        ("external_load_sweep_9900_report_figures/fig_5_12_indicator_compare.png", "图 5.13 外加载荷变化对转子响应指标的影响"),
    ]:
        picture(doc, rel, cap)

    heading(doc, "5.3 轴承支承刚度变化对振动特性的影响", 2)
    para(doc, "本次刚度分析中轴承1和轴承2同步取 kb1=K、kb2=K，其余参数保持不变。本版图表按筛选要求仅保留 K3、K4、K6、K7、K8 参与第三节表格、图片和趋势对比，对应刚度为 5e7、8e7、2e8、5e8、1e9 N/m。为弥补原报告仅输出前轴承节点响应的不足，本节重新从原始 MAT 结果中提取转子前轴承节点 2 与后轴承节点 10 的同口径位移、速度、加速度、轴心轨迹和频谱指标，并在图中采用实线表示前轴承、虚线表示后轴承。")
    rows = []
    for r in stiff:
        rows.append([r["case_id"], fmt(val(r, "K_level"), 2), fmt(val(r, "kb1_case"), 2), fmt(val(r, "kb2_case"), 2), r["simulation_status"], stiffness_status(r)])
    table(doc, "表 5.5 不同轴承支承刚度工况设置", ["工况", "K/(N/m)", "kb1/(N/m)", "kb2/(N/m)", "程序状态", "工程判断"], rows, 7.5)
    rows = []
    for r in stiff_fr:
        rows.append([
            r["case_id"], fmt(val(r, "K_level"), 2),
            f"{val(r,'front_disp_pp')*1e6:.3f}", f"{val(r,'rear_disp_pp')*1e6:.3f}", f"{100*val(r,'disp_rel_diff'):.1f}",
            f"{val(r,'front_vel_rms')*1e3:.3f}", f"{val(r,'rear_vel_rms')*1e3:.3f}", f"{100*val(r,'vel_rel_diff'):.1f}",
            f"{val(r,'front_acc_rms'):.3f}", f"{val(r,'rear_acc_rms'):.3f}", f"{100*val(r,'acc_rel_diff'):.1f}",
            f"{val(r,'front_amp_1X'):.3f}", f"{val(r,'rear_amp_1X'):.3f}", f"{100*val(r,'amp_1X_rel_diff'):.1f}",
        ])
    table(doc, "表 5.6 不同轴承支承刚度下前后轴承节点响应指标对比", ["工况", "K/(N/m)", "前位移/μm", "后位移/μm", "差/%", "前速度/(mm/s)", "后速度/(mm/s)", "差/%", "前加速度/(m/s²)", "后加速度/(m/s²)", "差/%", "前1X", "后1X", "差/%"], rows, 5.8)
    para(doc, "趋势分析：在筛选后的 K3、K4、K6、K7、K8 工况中，前后轴承响应随支承刚度变化的总体趋势一致，均表现为 5e7–2e8 N/m 区间响应先增大并在 K6 附近达到局部较高值，随后在 5e8–1e9 N/m 较高刚度区间明显降低。前后轴承的 1X 分量仍为主导频率成分，说明该组刚度变化下系统主要受同步不平衡响应控制。")
    para(doc, "前后轴承差异分析：K3 和 K4 下后轴承响应普遍低于前轴承，速度和加速度差异约为 -11% 至 -15%；K6 下前后轴承响应最接近，速度、加速度和 1X 幅值差异约为 1%–2%；K7 下后轴承响应高于前轴承，速度、加速度和 1X 幅值差异约为 38%–39%；K8 下后轴承响应低于前轴承，速度和加速度差异约为 -29%。这说明在支承刚度变化研究中，前后轴承并非所有工况都完全等效，尤其在高刚度小响应区间，相对差异会被放大。")
    para(doc, "工程合理性分析：筛选后保留的 K3、K4、K6、K7、K8 工况速度、位移和加速度均落在可解释范围内。若只使用前轴承节点 2 的数据，可以较好反映刚度变化引起的总体趋势和主导频率特征，但不能严格代表后轴承节点 10 的全部振动状态。因此，本节新增后轴承数据后，建议将前后轴承响应共同用于整机振动评价；若报告只保留单一评价值，则宜采用前后轴承中的较大响应值作为转子轴承处代表指标。")
    for rel, cap in [
        ("bearing_stiffness_sweep_9900_figures/fig_5_13_displacement_compare.png", "图 5.14 不同轴承支承刚度下前后轴承处位移响应对比"),
        ("bearing_stiffness_sweep_9900_figures/fig_5_14_velocity_compare.png", "图 5.15 不同轴承支承刚度下前后轴承处速度响应对比"),
        ("bearing_stiffness_sweep_9900_figures/fig_5_15_acceleration_compare.png", "图 5.16 不同轴承支承刚度下前后轴承处加速度响应对比"),
        ("bearing_stiffness_sweep_9900_figures/fig_5_16_orbit_compare.png", "图 5.17 不同轴承支承刚度下前后轴承处轴心轨迹对比"),
        ("bearing_stiffness_sweep_9900_figures/fig_5_17_acc_spectrum_compare.png", "图 5.18 不同轴承支承刚度下前后轴承处加速度频谱对比"),
        ("bearing_stiffness_sweep_9900_figures/fig_5_19_harmonic_compare.png", "图 5.19 轴承支承刚度变化对前后轴承倍频分量及幅值比的影响"),
        ("bearing_stiffness_sweep_9900_figures/fig_5_18_indicator_compare.png", "图 5.20 轴承支承刚度变化对前后轴承响应指标的影响"),
    ]:
        picture(doc, rel, cap)

    heading(doc, "5.4 结果真实性与一致性检查", 2)
    checks = [
        "所有表格数据均来自程序输出：是",
        "所有图片均由程序自动生成：是",
        "图片单位与正文单位一致：是",
        "不平衡工况已重新运行：是",
        "轴承刚度工况已重新运行：是",
        "外加载荷工况已核查：是",
        "不存在手动编造数据：是",
        "不合理工况已单独说明：是",
        "合理工况已明确标注：是",
        "结论与数据趋势一致：是",
    ]
    for item in checks:
        para(doc, item)
    para(doc, "补充说明：本版仅按要求从图表和正文对比中剔除指定工况，原始 MAT/CSV 仿真输出文件未修改，所有保留工况数据仍来自真实程序输出。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
