from pathlib import Path
import math

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook


SUMMARY = Path("bearing_stiffness_sweep_9900_summary.xlsx")
FIG_DIR = Path("bearing_stiffness_sweep_9900_figures")
OUT = Path("bearing_stiffness_sweep_9900_report_revised.docx")
MD_OUT = Path("bearing_stiffness_sweep_9900_report_revised.md")


def read_summary():
    wb = load_workbook(SUMMARY, data_only=True)
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        rows.append(dict(zip(headers, row)))
    return rows


def by_case(rows, case_id):
    for row in rows:
        if row.get("case_id") == case_id:
            return row
    raise KeyError(case_id)


def as_float(value):
    if value is None or value == "NaN":
        return float("nan")
    try:
        return float(value)
    except Exception:
        return float("nan")


def finite(value):
    return math.isfinite(as_float(value))


def fmt_sci(value, digits=3):
    if not finite(value):
        return "NaN"
    return f"{as_float(value):.{digits}e}"


def fmt_num(value, digits=5):
    if not finite(value):
        if value is None:
            return "NaN"
        return str(value)
    return f"{as_float(value):.{digits}g}"


def set_run_font(run, size=12, bold=False, name="SimSun"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, size=12, bold=False, align=None, after=6, line_spacing=1.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(after)
    if align is None:
        paragraph.paragraph_format.first_line_indent = Inches(0.29)
    else:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def set_cell_text(cell, text, bold=False, size=5.8):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, val in (("top", "60"), ("bottom", "60"), ("start", "80"), ("end", "80")):
        node = tc_mar.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            tc_mar.append(node)
        node.set(qn("w:w"), val)
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, top=None, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")
    if top:
        element = borders.find(qn("w:top"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(top))
        element.set(qn("w:color"), "000000")
    if bottom:
        element = borders.find(qn("w:bottom"))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(bottom))
        element.set(qn("w:color"), "000000")


def style_three_line_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=12, bottom=8)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=12)


def add_table(doc, caption, headers, rows, size=5.8):
    add_para(doc, caption, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=size)
    for data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(data):
            set_cell_text(cells[idx], value, size=size)
    style_three_line_table(table)
    add_para(doc, "", after=2, line_spacing=1.0)


def add_picture(doc, image_path, caption, width=8.2):
    if not image_path.exists():
        add_para(doc, f"{caption}（图片文件未生成）", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        return False
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_para(doc, caption, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    return True


def setting_rows(rows):
    unchanged = "除 kb1、kb2 外，其余参数保持原程序设置不变"
    out = []
    for row in rows:
        out.append([
            row.get("case_id"),
            fmt_sci(row.get("K_level"), 2),
            fmt_sci(row.get("kb1_case"), 2),
            fmt_sci(row.get("kb2_case"), 2),
            row.get("kb3_case"),
            row.get("Kzx_case"),
            row.get("Kzy_case"),
            fmt_num(row.get("rpm")),
            fmt_num(row.get("oneX_Hz")),
            fmt_num(row.get("unbalance_amplification")),
            fmt_sci(row.get("U1_used"), 3),
            fmt_sci(row.get("U2_used"), 3),
            fmt_sci(row.get("U3_used"), 3),
            fmt_sci(row.get("U4_used"), 3),
            fmt_sci(row.get("Famp1_1_used"), 3),
            fmt_sci(row.get("Famp1_2_used"), 3),
            fmt_sci(row.get("Famp1_3_used"), 3),
            fmt_sci(row.get("Famp1_4_used"), 3),
            unchanged,
            row.get("simulation_status"),
        ])
    return out


def response_rows(rows):
    out = []
    for row in rows:
        out.append([
            row.get("case_id"),
            fmt_sci(row.get("K_level"), 2),
            fmt_sci(row.get("disp_peak"), 3),
            fmt_sci(row.get("disp_pp"), 3),
            fmt_sci(row.get("vel_peak"), 3),
            fmt_sci(row.get("vel_rms"), 3),
            fmt_sci(row.get("acc_peak"), 3),
            fmt_sci(row.get("acc_rms"), 3),
            fmt_sci(row.get("orbit_radius_max"), 3),
            fmt_sci(row.get("amp_1X"), 3),
            fmt_sci(row.get("amp_2X"), 3),
            fmt_sci(row.get("amp_3X"), 3),
            fmt_num(row.get("ratio_2X_1X"), 4),
            fmt_num(row.get("ratio_3X_1X"), 4),
            row.get("simulation_status"),
        ])
    return out


def stable_trend_text(rows):
    k1, k2, k3, k4, k5 = [by_case(rows, c) for c in ("K1", "K2", "K3", "K4", "K5")]
    return [
        f"由表 5.6 可知，K1 和 K2 分别对应 1.0e6 N/m 和 1.0e7 N/m 的低支承刚度。二者的 simulation_status 均为 abnormal，位移峰峰值分别为 {fmt_sci(k1['disp_pp'])} m 和 {fmt_sci(k2['disp_pp'])} m，加速度 RMS 分别为 {fmt_sci(k1['acc_rms'])} m/s² 和 {fmt_sci(k2['acc_rms'])} m/s²，已经达到 10^-1 m 位移量级和 10^6 m/s² 加速度量级，远高于正常高压转子系统的小幅振动范围。这说明在极低支承刚度下，转子与机匣之间的耦合约束不足，系统出现异常大位移响应或数值发散趋势。因此，K1 和 K2 不能按稳定周期响应进行解释，也不参与稳定响应规律拟合。",
        f"对于完成全时长积分的 K3-K5 工况，随着轴承支承刚度由 1.0e8 N/m 增大到 1.0e10 N/m，位移峰峰值由 K3 的 {fmt_sci(k3['disp_pp'])} m 降至 K5 的 {fmt_sci(k5['disp_pp'])} m，速度 RMS 由 {fmt_sci(k3['vel_rms'])} m/s 降至 {fmt_sci(k5['vel_rms'])} m/s，加速度 RMS 由 {fmt_sci(k3['acc_rms'])} m/s² 降至 {fmt_sci(k5['acc_rms'])} m/s²，轴心轨迹最大半径由 {fmt_sci(k3['orbit_radius_max'])} m 降至 {fmt_sci(k5['orbit_radius_max'])} m。K4 的对应数值分别为位移峰峰值 {fmt_sci(k4['disp_pp'])} m、速度 RMS {fmt_sci(k4['vel_rms'])} m/s、加速度 RMS {fmt_sci(k4['acc_rms'])} m/s² 和轨迹最大半径 {fmt_sci(k4['orbit_radius_max'])} m，表明稳定工况范围内响应随刚度增大呈持续降低趋势。",
        "上述变化说明，在当前 9900 r/min 工况下，较高的轴承支承刚度增强了转子与机匣之间的约束作用，降低了转子在不平衡激励下的横向位移响应，同时也削弱了经轴承支承路径传递到机匣侧的振动能量。本次实际仿真中，K3-K5 的加速度 RMS 随支承刚度增大而降低，因此报告结论应以该真实结果为准，而不能写成刚度增大必然导致加速度增强。",
        f"机匣响应也支持这一判断。K3、K4、K5 的机匣节点加速度 RMS 分别为 {fmt_sci(k3['case_acc_rms'])} m/s²、{fmt_sci(k4['case_acc_rms'])} m/s² 和 {fmt_sci(k5['case_acc_rms'])} m/s²，整体随轴承支承刚度提高而降低。由于 K1、K2 已经处于 abnormal 状态，其机匣响应数值更多反映低刚度异常放大过程，不宜与 K3-K5 一起作为稳定传递规律拟合。"
    ]


def spectrum_text(rows):
    k1, k2, k3, k4, k5 = [by_case(rows, c) for c in ("K1", "K2", "K3", "K4", "K5")]
    return [
        f"频谱分析同样应区分 abnormal 工况和 completed 工况。对于完成全时长积分的 K3-K5，1X 幅值分别为 {fmt_sci(k3['amp_1X'])}、{fmt_sci(k4['amp_1X'])} 和 {fmt_sci(k5['amp_1X'])}；2X/1X 幅值比分别为 {fmt_num(k3['ratio_2X_1X'], 4)}、{fmt_num(k4['ratio_2X_1X'], 4)} 和 {fmt_num(k5['ratio_2X_1X'], 4)}；3X/1X 幅值比分别为 {fmt_num(k3['ratio_3X_1X'], 4)}、{fmt_num(k4['ratio_3X_1X'], 4)} 和 {fmt_num(k5['ratio_3X_1X'], 4)}。可以看出，在稳定支承刚度范围内，频谱中 1X 成分占主导，且 2X/1X、3X/1X 随支承刚度增大明显降低，说明系统主要表现为稳定同步响应。",
        f"K1 和 K2 虽然输出了频谱数据，1X 幅值分别达到 {fmt_sci(k1['amp_1X'])} 和 {fmt_sci(k2['amp_1X'])}，2X/1X 也接近 1，但由于其时域响应已经异常放大，这些频谱结果更多反映低刚度异常响应过程中的非稳态特征，不宜作为稳定频响规律进行解释。图 5.17 和图 5.19 中 K1、K2 的幅值远高于 K3-K5，主要用于反映低支承刚度下的异常放大现象。为了避免异常工况掩盖稳定工况规律，后续分析重点结合 K3-K5 的 completed 数据和图 5.20-图 5.24 进行。"
    ]


def conclusion_text(rows):
    k3, k5 = by_case(rows, "K3"), by_case(rows, "K5")
    return (
        f"综合来看，在 9900 r/min 固定转速下，轴承支承刚度对转子—轴承—机匣系统振动响应具有显著影响。K1 和 K2 工况对应 1.0e6-1.0e7 N/m 的低支承刚度范围，计算过程中出现异常放大，表明过低的轴承支承刚度会导致转子与机匣之间的耦合约束不足，系统难以形成稳定的小幅同步响应。对于完成全时长积分的 K3-K5 工况，随着轴承支承刚度由 1.0e8 N/m 增大至 1.0e10 N/m，转子位移峰峰值由 {fmt_sci(k3['disp_pp'])} m 降至 {fmt_sci(k5['disp_pp'])} m，速度 RMS 由 {fmt_sci(k3['vel_rms'])} m/s 降至 {fmt_sci(k5['vel_rms'])} m/s，加速度 RMS 由 {fmt_sci(k3['acc_rms'])} m/s² 降至 {fmt_sci(k5['acc_rms'])} m/s²，轴心轨迹最大半径由 {fmt_sci(k3['orbit_radius_max'])} m 降至 {fmt_sci(k5['orbit_radius_max'])} m，说明较高支承刚度增强了系统约束作用并抑制了横向振动响应。频谱结果表明，稳定工况下系统仍以 1X 同步响应为主，2X 和 3X 倍频成分占比较低，未出现明显高频冲击或倍频增强现象。因此，本组结果说明轴承支承刚度过低会引发异常放大，而在稳定支承刚度范围内，提高轴承支承刚度能够有效降低转子—轴承—机匣系统的振动响应。"
    )


def build_docx(rows):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)

    add_para(doc, "5.3 轴承支承刚度对转子—轴承—机匣系统振动响应的影响", size=16, bold=True, after=10, line_spacing=1.2)
    add_para(doc, "本研究固定转速为 9900 r/min，仅改变轴承支承刚度。轴承支承刚度采用 1.0e6、1.0e7、1.0e8、1.0e9 和 1.0e10 N/m 五组数量级工况，其余参数保持原程序设置不变，包括四个轮盘不平衡量、外加载荷、轴承阻尼、径向游隙、基础支承刚度、时间步长、积分方法和观测节点。批量计算前未放大不平衡激励量，unbalance_amplification=1；各工况开始时均重新设置 rpm=9900、wi=rpm*2*pi/60、data(5)=0，并调用原 Newmark-Newton 程序完成时域积分。")
    add_para(doc, "本节采用五个数量级的轴承支承刚度工况，目的是考察支承刚度从弱约束到强约束变化时系统响应的敏感性。需要说明的是，1.0e6 和 1.0e7 N/m 对当前高压转子—轴承—机匣耦合系统而言属于较低支承刚度范围，其计算结果主要用于判断支承约束不足时系统是否发生异常放大；而 1.0e8-1.0e10 N/m 工况完成了全时长积分，更适合用于讨论稳定支承范围内响应随刚度变化的规律。因此，后续规律分析以 K3-K5 的 completed 数据为主，K1-K2 仅用于说明低刚度异常放大现象。")
    add_para(doc, "程序结构检查表明，当前耦合模型中轴承支承刚度的实际入口位于 mian.m 的转子—机匣组合矩阵装配阶段。程序将转子轴承节点 r1=2、r2=10 与机匣轴承座节点 c1=3、c2=10 通过 kb1、kb2 连接，并在 x、y 两个平动方向将 +kb、-kb、-kb、+kb 形式的耦合刚度直接加入整体刚度矩阵 KK。因此，本节实际修改的是 kb1 和 kb2，并令二者同时等于当前 K_level。程序中未发现第三个轴承支承刚度 kb3，也未发现以 K_b 命名的轴承支承刚度矩阵；K_D.m 中的 Kzx、Kzy 只在非耦合固定支承模型中使用，而当前 data(8)=2 的转子—机匣耦合模型下该通道不作为轴承支承刚度入口。K_D_case.m 中的 Kzx、Kzy 对应机匣基础支承刚度，按单因素要求保持不变。")
    add_para(doc, "F_bearing.m 中的 C_b=data(6) 是滚动轴承赫兹接触刚度，参与非线性轴承接触力计算；data(9) 为径向游隙，data(10) 为外加载荷入口。本次修订不改变仿真原始数据，也不调整任何工况参数，报告表格和趋势判断仍全部来自 bearing_stiffness_sweep_9900_summary.xlsx 以及对应 MAT 文件。")

    add_table(doc, "表 5.5 不同轴承支承刚度工况设置",
              ["case_id", "K_level", "kb1_case", "kb2_case", "kb3_case", "Kzx_case", "Kzy_case", "rpm", "1X", "unbalance_amp", "U1_used", "U2_used", "U3_used", "U4_used", "Famp1_1", "Famp1_2", "Famp1_3", "Famp1_4", "其他参数保持不变说明", "status"],
              setting_rows(rows), size=4.6)
    add_table(doc, "表 5.6 不同轴承支承刚度下系统振动响应指标",
              ["case_id", "K_level", "disp_peak", "disp_pp", "vel_peak", "vel_rms", "acc_peak", "acc_rms", "orbit_R", "amp_1X", "amp_2X", "amp_3X", "2X/1X", "3X/1X", "status"],
              response_rows(rows), size=5.9)
    add_para(doc, "表 5.6 中 K1 和 K2 虽然有数值输出，但 simulation_status 均为 abnormal；其位移、速度和加速度结果反映的是低支承刚度下的异常放大或数值失稳趋势，不能作为正常稳定周期响应进行外推。K3、K4 和 K5 的状态为 completed，是后续稳定响应规律分析的主要依据。")

    original_figs = [
        ("fig_5_13_displacement_compare.png", "图 5.13 不同轴承支承刚度下转子节点位移响应对比"),
        ("fig_5_14_velocity_compare.png", "图 5.14 不同轴承支承刚度下转子节点速度响应对比"),
        ("fig_5_15_acceleration_compare.png", "图 5.15 不同轴承支承刚度下转子节点加速度响应对比"),
        ("fig_5_16_orbit_compare.png", "图 5.16 不同轴承支承刚度下轴心轨迹对比"),
        ("fig_5_17_acc_spectrum_compare.png", "图 5.17 不同轴承支承刚度下加速度频谱对比"),
        ("fig_5_18_indicator_compare.png", "图 5.18 轴承支承刚度变化对主要响应指标的影响"),
        ("fig_5_19_harmonic_compare.png", "图 5.19 轴承支承刚度变化对倍频分量及幅值比的影响"),
    ]
    for name, caption in original_figs:
        add_picture(doc, FIG_DIR / name, caption)
    add_para(doc, "图 5.13-图 5.19 保留了 K1-K5 全工况对比。由于低刚度工况 K1、K2 响应幅值远高于 K3-K5，原全工况图主要用于显示低支承刚度下的异常放大现象；稳定工况趋势需结合表 5.6 中 K3-K5 的 completed 数据以及下列局部放大图进行分析。")

    zoom_figs = [
        ("fig_5_20_completed_displacement_zoom.png", "图 5.20 稳定工况 K3-K5 下转子节点位移响应局部放大图"),
        ("fig_5_21_completed_velocity_zoom.png", "图 5.21 稳定工况 K3-K5 下转子节点速度响应局部放大图"),
        ("fig_5_22_completed_acceleration_zoom.png", "图 5.22 稳定工况 K3-K5 下转子节点加速度响应局部放大图"),
        ("fig_5_23_completed_orbit_zoom.png", "图 5.23 稳定工况 K3-K5 下轴心轨迹局部放大图"),
        ("fig_5_24_completed_acc_spectrum_zoom.png", "图 5.24 稳定工况 K3-K5 下加速度频谱局部放大图"),
    ]
    missing_zoom = []
    for name, caption in zoom_figs:
        ok = add_picture(doc, FIG_DIR / name, caption)
        if not ok:
            missing_zoom.append(name)
    if missing_zoom:
        add_para(doc, "K3-K5 局部放大图生成失败或文件缺失：" + "、".join(missing_zoom) + "。因此稳定工况趋势主要结合表 5.6 中 K3-K5 的数值进行分析。")

    for paragraph in stable_trend_text(rows):
        add_para(doc, paragraph)
    for paragraph in spectrum_text(rows):
        add_para(doc, paragraph)
    add_para(doc, conclusion_text(rows))
    doc.save(OUT)


def md_table(headers, table_rows):
    lines = ["|" + "|".join(headers) + "|", "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in table_rows:
        lines.append("|" + "|".join(str(v) for v in row) + "|")
    return lines


def write_markdown(rows):
    lines = [
        "# 5.3 轴承支承刚度对转子—轴承—机匣系统振动响应的影响",
        "",
        "本研究固定转速为 9900 r/min，仅改变 `mian.m` 中实际装配进整体刚度矩阵 `KK` 的 `kb1` 和 `kb2`。主要稳定趋势基于 K3-K5 completed 工况，K1-K2 仅用于说明低刚度异常放大现象。",
        "",
        "表 5.5 不同轴承支承刚度工况设置",
    ]
    setting_headers = ["case_id", "K_level", "kb1_case", "kb2_case", "kb3_case", "Kzx_case", "Kzy_case", "rpm", "1X", "unbalance_amp", "U1_used", "U2_used", "U3_used", "U4_used", "Famp1_1", "Famp1_2", "Famp1_3", "Famp1_4", "其他参数保持不变说明", "status"]
    lines.extend(md_table(setting_headers, setting_rows(rows)))
    lines += ["", "表 5.6 不同轴承支承刚度下系统振动响应指标"]
    response_headers = ["case_id", "K_level", "disp_peak", "disp_pp", "vel_peak", "vel_rms", "acc_peak", "acc_rms", "orbit_R", "amp_1X", "amp_2X", "amp_3X", "2X/1X", "3X/1X", "status"]
    lines.extend(md_table(response_headers, response_rows(rows)))
    lines.append("")
    lines.append("K1 和 K2 虽然有数值输出，但状态为 abnormal；K3-K5 为 completed，是稳定响应分析的主要依据。")
    lines.append("")
    figs = [
        ("fig_5_13_displacement_compare.png", "图 5.13 不同轴承支承刚度下转子节点位移响应对比"),
        ("fig_5_14_velocity_compare.png", "图 5.14 不同轴承支承刚度下转子节点速度响应对比"),
        ("fig_5_15_acceleration_compare.png", "图 5.15 不同轴承支承刚度下转子节点加速度响应对比"),
        ("fig_5_16_orbit_compare.png", "图 5.16 不同轴承支承刚度下轴心轨迹对比"),
        ("fig_5_17_acc_spectrum_compare.png", "图 5.17 不同轴承支承刚度下加速度频谱对比"),
        ("fig_5_18_indicator_compare.png", "图 5.18 轴承支承刚度变化对主要响应指标的影响"),
        ("fig_5_19_harmonic_compare.png", "图 5.19 轴承支承刚度变化对倍频分量及幅值比的影响"),
        ("fig_5_20_completed_displacement_zoom.png", "图 5.20 稳定工况 K3-K5 下转子节点位移响应局部放大图"),
        ("fig_5_21_completed_velocity_zoom.png", "图 5.21 稳定工况 K3-K5 下转子节点速度响应局部放大图"),
        ("fig_5_22_completed_acceleration_zoom.png", "图 5.22 稳定工况 K3-K5 下转子节点加速度响应局部放大图"),
        ("fig_5_23_completed_orbit_zoom.png", "图 5.23 稳定工况 K3-K5 下轴心轨迹局部放大图"),
        ("fig_5_24_completed_acc_spectrum_zoom.png", "图 5.24 稳定工况 K3-K5 下加速度频谱局部放大图"),
    ]
    for name, caption in figs:
        lines.append(f"![{caption}]({(FIG_DIR / name).resolve().as_posix()})")
        lines.append("")
    lines.extend(stable_trend_text(rows))
    lines.append("")
    lines.extend(spectrum_text(rows))
    lines.append("")
    lines.append(conclusion_text(rows))
    lines.append("")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def verify_inputs():
    required = [
        SUMMARY,
        Path("bearing_stiffness_sweep_9900_report.docx"),
        Path("results_bearing_stiffness_K1.mat"),
        Path("results_bearing_stiffness_K2.mat"),
        Path("results_bearing_stiffness_K3.mat"),
        Path("results_bearing_stiffness_K4.mat"),
        Path("results_bearing_stiffness_K5.mat"),
    ]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise SystemExit("Missing required input files: " + ", ".join(missing))


def build():
    verify_inputs()
    rows = read_summary()
    write_markdown(rows)
    build_docx(rows)
    print(f"Generated {OUT} and {MD_OUT}")


if __name__ == "__main__":
    build()
